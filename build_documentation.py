from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("artifacts/documentation")
OUT_DIR.mkdir(parents=True, exist_ok=True)

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F7FB"
BORDER = "C9D3DF"
WHITE = "FFFFFF"
BLACK = "111827"
CAUTION = "7A5A00"
CAUTION_FILL = "FFF7D6"
RED = "9B1C1C"
RED_FILL = "FDECEC"
GREEN = "1F5E3B"
GREEN_FILL = "EAF6EF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, header=True, fill=LIGHT_BLUE):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        if len(row.cells) != len(widths_dxa):
            continue
        if row_index == 0 and header:
            tr_pr = row._tr.get_or_add_trPr()
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0 and header:
                set_cell_shading(cell, fill)


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def configure_styles(doc, preset):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10 if preset == "business" else 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 16 if preset == "business" else 18, 8 if preset == "business" else 10),
        "Heading 2": (13, BLUE, 12 if preset == "business" else 14, 6 if preset == "business" else 7),
        "Heading 3": (12, DARK_BLUE, 8 if preset == "business" else 10, 4 if preset == "business" else 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True

    for style_name, size, color in (("Table Text", 9.5, BLACK), ("Table Header", 9.5, NAVY), ("Code Text", 9, NAVY)):
        if style_name in styles:
            style = styles[style_name]
        else:
            style = styles.add_style(style_name, 1)
        style.font.name = "Calibri" if style_name != "Code Text" else "Consolas"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    return section


def add_page_field(paragraph):
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_header_footer(doc, header_text, footer_text):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, after=3, line=1.0)
    run = p.add_run(header_text)
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(p, before=3, after=0, line=1.0)
    run = p.add_run(footer_text + " | ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(p)


def add_title_block(doc, kicker, title, subtitle, metadata):
    p = doc.add_paragraph()
    set_para_spacing(p, before=6, after=9, line=1.0)
    run = p.add_run(kicker.upper())
    set_run_font(run, size=9.5, color=BLUE, bold=True)

    p = doc.add_paragraph()
    set_para_spacing(p, after=5, line=1.0)
    run = p.add_run(title)
    set_run_font(run, size=26, color=NAVY, bold=True)

    p = doc.add_paragraph()
    set_para_spacing(p, after=13, line=1.10)
    run = p.add_run(subtitle)
    set_run_font(run, size=14, color=MUTED)

    for label, value in metadata:
        p = doc.add_paragraph()
        set_para_spacing(p, after=2, line=1.0)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=BLACK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color=BLACK)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_para_spacing(p, after=6, line=1.10)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_note(doc, label, text, fill=PALE_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA], header=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=0, line=1.15)
    label_run = p.add_run(label + " ")
    set_run_font(label_run, size=10.5, color=color, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.5, color=BLACK)
    spacer = doc.add_paragraph()
    set_para_spacing(spacer, after=2, line=1.0)


def add_bullet_numbering(doc, kind="bullet", left=540, hanging=270):
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abstract or [0]) + 1
    num_id = max(existing_num or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_list_item(doc, text, num_id, after=4, line=1.167):
    p = doc.add_paragraph()
    set_para_spacing(p, after=after, line=line)
    apply_numbering(p, num_id)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_code(doc, lines):
    p = doc.add_paragraph(style="Code Text")
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF3F8")
    p_pr.append(shd)
    for index, line in enumerate(lines.splitlines()):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=9, color=NAVY)
    return p


def add_table(doc, headers, rows, widths, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, header=True)
    for cell, header in zip(table.rows[0].cells, headers):
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Header"]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        set_run_font(run, size=9.5, color=NAVY, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cell = cells[index]
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            if alignments and index < len(alignments):
                p.alignment = alignments[index]
            run = p.add_run(str(value))
            set_run_font(run, size=9.5, color=BLACK)
    set_table_geometry(table, widths, header=True)
    spacer = doc.add_paragraph()
    set_para_spacing(spacer, after=2, line=1.0)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p


def add_source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=8.5, color=MUTED, italic=True)


def set_core_properties(doc, title, subject):
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "ASI Spirits"
    props.keywords = "TSYS, PAX, batch reporting, MXConnect, operations"


def build_technical_manual():
    doc = Document()
    configure_styles(doc, "business")
    section = configure_page(doc)
    configure_header_footer(doc, "TSYS_PAX_BATCH_REPORT | Technical Master Manual", "ASI Spirits")
    set_core_properties(doc, "TSYS_PAX_BATCH_REPORT - Technical Master Manual", "Technical reference for setup, operation, API behavior, outputs, and maintenance")

    add_title_block(
        doc,
        "Technical Master Manual",
        "TSYS_PAX_BATCH_REPORT",
        "Store-level TSYS/PAX batch exception reporting for Windows",
        [
            ("Documentation baseline", "July 21, 2026"),
            ("Implementation", "get_tsys_batch.py packaged as UI and CLI executables"),
            ("Repository", "asispirits/GET_TSYS_BATCH"),
            ("Primary output", "pinpad_batch_not_closed_<N>_days.csv"),
            ("API key", "Embedded in official Windows builds; environment fallback for development"),
        ],
    )
    add_note(
        doc,
        "Core operating rule.",
        "The utility reports active TSYS stores with approved authorization activity in the configured authorization window and no accepted TSYS batch for the store account in the selected batch window. It is intentionally store-level and does not claim that a particular terminal or physical device failed to batch.",
        fill=CAUTION_FILL,
        color=CAUTION,
    )

    add_heading(doc, "1. Purpose and scope", 1)
    add_body(doc, "TSYS_PAX_BATCH_REPORT replaces the earlier multi-script comparison process with one packaged utility. It combines the TSYS batch export, the active merchant roster, and authorization activity into a repeatable store-level exception report.")
    add_body(doc, "The program is designed for Windows desktop use through a minimal Tkinter UI and for unattended execution through the console executable and Windows Task Scheduler. The source remains a single Python program; the Windows workflow creates separate UI and CLI entry points from that source.")
    add_body(doc, "The program does not use serial numbers, vnumber, external-terminal files, or manual terminal-to-device crosswalks as part of the alert decision. Optional configuration rows are display overrides only.")

    add_heading(doc, "2. What the program produces", 1)
    add_table(
        doc,
        ["File", "Purpose", "Alert-safe interpretation"],
        [
            ("pinpad_batch_not_closed_<N>_days.csv", "Primary store-level exception report.", "Rows are stores/accounts that passed the active, approved-authorization, and no-accepted-batch rules."),
            ("in_use_not_batched_detail.csv", "Terminal-level authorization detail beneath each flagged account.", "Shows terminalNumber returned with authorization activity; it is supporting detail, not proof that that terminal failed to batch."),
            ("needs_mapping_or_review.csv", "Rows excluded from the primary report because required store display data is not safe to use.", "Review-only output. It should not be emailed as an exception list."),
            ("batch_history.csv", "Accepted batch records returned for the current report window.", "Audit trail for the batch data used by the run."),
            ("termid_account_history.csv", "Distinct accountNumber + termID pairs and their latest batch date in the returned window.", "Historical reference for termID/account analysis; not used to assert a device identity."),
            ("tsys_batch_report_summary_<N>_days.docx", "One-page summary of the CSV outputs created by the current run.", "Convenience summary; the CSV files remain the detailed source outputs."),
        ],
        [2050, 3000, 4310],
    )
    add_body(doc, "If a raw batch path is supplied with -tf/--tsys_filename, the utility also writes the unfiltered batch export to that path. The default report run does not create that optional compatibility file.")
    add_body(doc, "Each normal report run creates a new timestamped subfolder under outputDirectory using YYYYMMDD_HHMMSS. If another run starts in the same second, a numeric suffix such as _01 is added. The CSV files and DOCX summary for the run are kept together in that folder. Historical refreshes use the same pattern under outputDirectory/historical.")

    add_heading(doc, "3. System components", 1)
    add_table(
        doc,
        ["Component", "Responsibility"],
        [
            ("TSYS_PAX_BATCH_REPORT.exe", "Windowed UI for selecting the config file, output folder, report timeframe, and report actions."),
            ("TSYS_PAX_BATCH_REPORT_CLI.exe", "Console entry point used by the UI and by Task Scheduler."),
            ("config.json", "Portable JSON settings and optional account-level display overrides."),
            ("MXConnect API", "Authentication, TSYS batch export, active merchant roster, and authorization activity."),
            ("GitHub Actions workflow", "Injects the repository secret at build time, builds both Windows executables, and packages them with config.json."),
        ],
        [2700, 6660],
    )
    add_body(doc, "The UI starts the sibling CLI executable when the program is frozen into an .exe. Both executables therefore need to remain in the same folder for the Run report and Refresh historical data buttons to work correctly.")

    add_heading(doc, "4. MXConnect API contract", 1)
    add_table(
        doc,
        ["Operation", "Method and path", "Use"],
        [
            ("Authenticate", "POST /security/v1/apiKey/authenticate", "Sends {\"value\": API_KEY}; expects a token."),
            ("TSYS batch export", "POST /report/v1/tsys/batch/export", "Fetches batch records for the selected date window using paginated scroll responses."),
            ("Authorization export", "POST /report/v1/tsys/authorization/export", "Fetches authorization detail for candidate account numbers using dr_type=q and the configured dr_quick window."),
            ("Active merchant roster", "GET /boarding/v1/uar", "Fetches the merchant roster in pages and supplies the authoritative active TSYS store list."),
        ],
        [2100, 3550, 3710],
    )
    add_body(doc, "The configured base URL is https://api.mxconnect.com. If an older config contains the authentication path at the end of apiBaseUrl, the code removes that suffix before building the other endpoint URLs.")
    add_body(doc, "Pagination is fail-closed. If MXConnect reports more records but omits a scroll ID, or if the returned count does not match the reported total, the run raises an error rather than producing a partial report.")

    add_heading(doc, "5. Report decision logic", 1)
    add_body(doc, "The decision sequence is deliberately account/store based. The primary report is not a claim that a specific terminal or device was the one that failed to batch.")
    num = add_bullet_numbering(doc, "decimal", left=540, hanging=270)
    for item in [
        "Read config.json and resolve the output directory relative to the config file when the configured path is relative.",
        "Use the API key embedded in the official Windows executable. When the embedded key is unavailable, fall back to the environment variable named by apiKeyEnvironmentVariable; the default name is MXCONNECT_API_KEY.",
        "Fetch the TSYS batch export for the selected batch lookback. A record is accepted when rejected is one of no, false, 0, or n. Rejected records do not count as a successful batch.",
        "Fetch the active merchant roster and keep only product ID 3, active=true, and account statuses other than Closed, Terminated, or Suspended.",
        "Build the candidate set as active TSYS accounts minus accounts with at least one accepted batch record in the report window.",
        "Fetch authorization detail only for candidate accounts. When requireAuthorizationActivity is enabled, only records with authorizationResponseStatus=approved are included.",
        "Aggregate approved authorized amount and count by account and terminalNumber, then roll the amounts up to the store/account for the primary CSV.",
        "Use the active roster location name as the authoritative store name. Apply a configured URL, store-name, or device value only as an optional display override for a known accountNumber.",
        "Write the primary row only when the account is in the active roster, the store name is available, and there is no conflicting configured display override.",
    ]:
        add_list_item(doc, item, num, after=5, line=1.167)
    add_note(doc, "Important.", "The primary CSV leaves TERMID blank by design. The current reliable evidence supports store/account exception reporting, not an exact terminal-to-batch failure claim.", fill=RED_FILL, color=RED)

    add_heading(doc, "6. Time windows and amount semantics", 1)
    add_body(doc, "batchLookbackDays controls the batch lookback used by the primary report. The date-window helper calculates a local cutoff at 04:00 and passes the resulting inclusive start and end dates to the TSYS batch export. The report UI exposes this as Report timeframe (days), with a range of 1 through 365.")
    add_body(doc, "authorizationWindow controls the MXConnect quick authorization window. The packaged template uses last_24_h. The UI keeps this option out of the main screen; it can be changed in config.json or overridden on the CLI with --auth-window.")
    add_body(doc, "AMOUNT in the primary CSV is the sum of approved authorizedAmount values returned for the candidate account during the authorization window. It is an exposure/supporting amount, not the TSYS batch total and not a terminal-specific batch amount.")

    add_heading(doc, "7. Configuration reference", 1)
    add_code(doc, '''{
  "apiBaseUrl": "https://api.mxconnect.com",
  "apiKeyEnvironmentVariable": "MXCONNECT_API_KEY",
  "batchLookbackDays": 3,
  "historicalLookbackDays": 90,
  "authorizationWindow": "last_24_h",
  "requireAuthorizationActivity": true,
  "outputDirectory": "./tsys-auditdata",
  "devices": []
}''')
    add_table(
        doc,
        ["Key", "Default", "Technical meaning"],
        [
            ("apiBaseUrl", "https://api.mxconnect.com", "MXConnect base URL; authentication and report paths are appended by the program."),
            ("apiKeyEnvironmentVariable", "MXCONNECT_API_KEY", "Fallback environment variable name for source or locally built executables. Official Windows builds use the embedded build-time key first."),
            ("batchLookbackDays", "3", "Primary batch report lookback. UI changes are saved here."),
            ("historicalLookbackDays", "90", "Lookback for the separate historical refresh action."),
            ("authorizationWindow", "last_24_h", "MXConnect quick authorization window for candidate accounts."),
            ("requireAuthorizationActivity", "true", "The current report requires approved authorization activity; the UI enforces true."),
            ("outputDirectory", "./tsys-auditdata", "Output folder. Relative paths are resolved beside config.json."),
            ("devices", "[]", "Optional account-level display override rows. They are not device validation records."),
        ],
        [2500, 1850, 5110],
    )
    add_body(doc, "Optional devices rows may contain enabled, url, storeName, device, and accountNumber. An enabled row with an accountNumber can override URL, store name, and device text for display. Rows without accountNumber are ignored by the report logic. Conflicting enabled rows for the same account are sent to needs_mapping_or_review.csv and excluded from the primary CSV.")

    add_heading(doc, "8. UI operation", 1)
    num = add_bullet_numbering(doc, "decimal", left=540, hanging=270)
    for item in [
        "Place TSYS_PAX_BATCH_REPORT.exe, TSYS_PAX_BATCH_REPORT_CLI.exe, and config.json in the same folder.",
        "For the official Windows package, no API-key setup is required. The key is embedded during the GitHub Actions build. If running the Python source or an executable built without the repository secret, set MXCONNECT_API_KEY in the Windows User or Machine environment and restart the program after changing it.",
        "Open TSYS_PAX_BATCH_REPORT.exe. Use Open config if config.json is not beside the program, or use the displayed path.",
        "Choose the output folder and set Report timeframe (days). The default is 3.",
        "Click Save config before running. The UI stores the output folder as a portable relative path when it is under the config folder.",
        "Click Run report. Read the log panel for authentication, fetch progress, counts, and output paths.",
        "Use Refresh historical data only when an updated historical batch/roster snapshot is needed; it writes into a timestamped folder under the historical subfolder.",
    ]:
        add_list_item(doc, item, num, after=5, line=1.167)

    add_heading(doc, "9. Command-line and Task Scheduler operation", 1)
    add_body(doc, "The CLI is intended for unattended execution. It uses the same config and API-key selection as the UI: the official Windows build uses its embedded key, while source or local fallback builds can use MXCONNECT_API_KEY.")
    add_code(doc, '''TSYS_PAX_BATCH_REPORT_CLI.exe --run-report --config "C:\\TSYS_PAX_BATCH_REPORT\\config.json"''')
    add_body(doc, "Useful command-line options:")
    add_table(
        doc,
        ["Option", "Purpose"],
        [
            ("--run-report", "Run the primary report without opening the UI."),
            ("--refresh-historical", "Refresh historical batch, termID/account, and active roster CSV files."),
            ("--config PATH", "Use a specific config.json path."),
            ("--batch-days N", "Override batchLookbackDays for one run."),
            ("--historical-days N", "Override historicalLookbackDays for one refresh."),
            ("--auth-window VALUE", "Override authorizationWindow for one run."),
            ("--output PATH", "Override the primary email CSV path."),
            ("--review-output PATH", "Override the review CSV path."),
            ("--tsys_filename PATH", "Also write the raw unfiltered batch export for compatibility."),
        ],
        [2700, 6660],
    )
    add_body(doc, "For Task Scheduler, set Start in to the folder containing the executables, use the CLI executable as the program, and include --run-report --config followed by the full config path in the arguments. With the official Windows package, the scheduled account needs access to the executables, config file, and output folder; it does not need an API-key environment variable.")

    add_heading(doc, "10. Output field reference", 1)
    add_table(
        doc,
        ["File / field", "Meaning", "Use / caution"],
        [
            ("Primary / URL", "Configured display URL override when present.", "May be blank. It is not obtained from the active roster."),
            ("Primary / STORENAME", "Configured store-name override or active roster location name.", "Authoritative display label for the store row."),
            ("Primary / DEVICE", "Configured display device or PAX default.", "Display text only; not a device identity assertion."),
            ("Primary / AMOUNT", "Sum of approved authorizedAmount values for the account.", "Exposure/supporting amount, not batch amount."),
            ("Primary / TERMID", "Blank in the current implementation.", "Blank intentionally to prevent unsupported terminal claims."),
            ("Detail / terminalNumber", "Terminal number returned in authorization detail.", "Supporting activity detail; not proof of unbatched terminal."),
            ("History / termID", "termID returned by accepted batch records.", "Used for historical account/term analysis only."),
        ],
        [2500, 3500, 3360],
    )

    add_heading(doc, "11. Security and deployment controls", 1)
    for item in [
        "Never place the API key in config.json, the source repository, a CSV export, or a screenshot. The GitHub Actions repository secret is the build input; it is not committed.",
        "The official Windows executables contain the API key so end users do not need PowerShell setup. Any secret embedded in a distributed executable can potentially be extracted; use a restricted key and rotate it if distribution expands.",
        "Use the MXCONNECT_API_KEY environment variable only for source execution, local builds without an embedded key, or emergency fallback. Restart the UI or task process after changing environment variables.",
        "Keep the output directory access-controlled because reports contain merchant names, account-linked activity, and authorized amounts.",
        "Do not treat a CSV row as proof of a specific physical device failure. The report is intentionally conservative at the store/account level.",
        "Keep the UI and CLI executables from the same build together with the config file. The UI depends on the sibling CLI executable for Run report and Refresh historical data.",
    ]:
        add_list_item(doc, item, add_bullet_numbering(doc, "bullet", left=540, hanging=270), after=5, line=1.167)

    add_heading(doc, "12. GitHub Actions build and release", 1)
    add_body(doc, "The repository workflow is named Build TSYS_PAX_BATCH_REPORT for Windows. It runs on pushes that change the source, requirements, template config, or workflow, and can also be started manually with workflow_dispatch. The repository must have an Actions secret named MXCONNECT_API_KEY.")
    num = add_bullet_numbering(doc, "decimal", left=540, hanging=270)
    for item in [
        "Checkout the repository and open the Actions tab.",
        "Select Build TSYS_PAX_BATCH_REPORT for Windows and use Run workflow against main for a manual build.",
        "The workflow installs requests and PyInstaller, creates a temporary embedded_api_key.py module from the MXCONNECT_API_KEY repository secret, verifies Python compilation, and builds one windowed and one console executable.",
        "The temporary API-key source module is removed after both executables are built. The key is not included in config.json or the repository checkout.",
        "The workflow copies config.template.json to dist/config.json and packages the three files as TSYS_PAX_BATCH_REPORT-windows.zip.",
        "Download the artifact from a successful run. The artifact is not committed to the repository.",
    ]:
        add_list_item(doc, item, num, after=5, line=1.167)
    add_note(doc, "Build limitation.", "The workflow produces Windows executables on windows-latest. macOS development can validate the Python source but does not replace the Windows packaging run.", fill=CAUTION_FILL, color=CAUTION)

    add_heading(doc, "13. Failure handling and troubleshooting", 1)
    add_table(
        doc,
        ["Message / symptom", "Likely cause", "Corrective action"],
        [
            ("Set MXCONNECT_API_KEY before running the API report.", "The program is running from source or from an executable built without an embedded key.", "Use the official Windows artifact, or set the fallback variable for the same Windows user that runs the UI/task. Never print the key."),
            ("MXConnect authentication failed with 404.", "Wrong base URL, stale config, or endpoint mismatch.", "Use https://api.mxconnect.com as apiBaseUrl and confirm the build came from the current repository workflow."),
            ("Report exited with code 1.", "The CLI child reported a runtime/config/API error.", "Read the preceding log line; correct the root error before relying on the CSV."),
            ("No primary rows.", "No approved authorization activity among candidate stores, or all candidate accounts had accepted batches.", "Check batch window, authorization window, active roster, and the detail/review files."),
            ("Rows appear in needs_mapping_or_review.csv.", "Active roster display data is missing or config overrides conflict.", "Correct the store name or duplicate account override; do not copy review rows into the email list."),
            ("Output folder appears not to save.", "The UI was not saving the active config path or the process was using a different config.", "Confirm the Config file path shown in the UI, click Save config, and inspect that exact file."),
        ],
        [2800, 3000, 3560],
    )
    add_body(doc, "The program clears the primary email CSV before starting an API run. If a run fails, an old primary CSV is not left looking like the current result. Always confirm the run completed successfully before distributing an output file.")

    add_heading(doc, "14. Known limitations and interpretation", 1)
    for item in [
        "The report detects account/store-level exceptions. It does not identify the exact device that did not batch.",
        "TERMID is blank in the primary output by design. The detail file may contain terminalNumber from authorization activity, but that field is not joined to a batch failure claim.",
        "The batch export may contain termID and accountNumber, but the current alert decision uses the presence of any accepted batch for the account in the selected window.",
        "URL is optional display data maintained in config.json; it is not required for the API decision.",
        "The historical files represent the records returned for the chosen lookback and are useful for review, not a complete permanent TSYS ledger.",
        "API availability, authorization-window semantics, and roster freshness remain external dependencies.",
    ]:
        add_list_item(doc, item, add_bullet_numbering(doc, "bullet", left=540, hanging=270), after=5, line=1.167)

    add_heading(doc, "15. Validation and change control", 1)
    add_body(doc, "The current repository build was validated with Python compilation, workflow YAML validation, source-level synthetic store/batch rule checks, and live authentication/batch API checks during development. A complete live end-to-end report run depends on API response time and merchant-roster volume.")
    add_body(doc, "When changing the report logic, update the source, run the targeted validation, build through GitHub Actions, and compare the primary CSV, detail CSV, review CSV, and batch history for a known test window. Keep a copy of the prior executable and config until the new build has been accepted.")

    add_heading(doc, "16. Rollback", 1)
    add_body(doc, "Rollback is file-based. Stop scheduled execution, restore the prior pair of executables and the prior config.json, and rerun the prior executable for the required window. Official packages use their embedded build-time key; source or locally built fallback packages may still require MXCONNECT_API_KEY.")

    add_source_note(doc, "Source of truth for this manual: get_tsys_batch.py, config.template.json, requirements.txt, and .github/workflows/build-windows.yml in asispirits/GET_TSYS_BATCH.")
    path = OUT_DIR / "TSYS_PAX_BATCH_REPORT_Technical_Master_Manual.docx"
    doc.save(path)
    return path


def build_user_guide():
    doc = Document()
    configure_styles(doc, "compact")
    configure_page(doc)
    configure_header_footer(doc, "TSYS_PAX_BATCH_REPORT | User Guide", "ASI Spirits")
    set_core_properties(doc, "TSYS_PAX_BATCH_REPORT - User Setup and Operation Guide", "Short operator guide for setup, daily use, and output interpretation")

    add_title_block(
        doc,
        "User Setup and Daily Operation Guide",
        "TSYS_PAX_BATCH_REPORT",
        "How to install, configure, run, and interpret the store-level batch report",
        [
            ("Audience", "Operations, support, and report users"),
            ("Platform", "Windows"),
            ("Default report", "3-day batch lookback with last_24_h authorization activity"),
        ],
    )
    add_note(doc, "What it tells you.", "The report identifies active TSYS stores with approved authorization activity but no accepted batch for the store account in the selected timeframe. It reports stores, not a verified individual terminal failure.", fill=CAUTION_FILL, color=CAUTION)

    add_heading(doc, "1. Put the files in one folder", 1)
    add_body(doc, "Keep these three files together in the same folder:")
    bullet = add_bullet_numbering(doc, "bullet", left=540, hanging=270)
    for item in [
        "TSYS_PAX_BATCH_REPORT.exe - the setup and manual-run UI.",
        "TSYS_PAX_BATCH_REPORT_CLI.exe - the unattended/Task Scheduler runner.",
        "config.json - the settings file saved by the UI.",
    ]:
        add_list_item(doc, item, bullet, after=4, line=1.25)
    add_body(doc, "Do not rename the CLI executable if you plan to use the UI Run report button. The UI looks for the sibling file named TSYS_PAX_BATCH_REPORT_CLI.exe.")

    add_heading(doc, "2. API key setup", 1)
    add_body(doc, "The official Windows package already contains the build-time API key. End users do not need to set an environment variable or run PowerShell. The key is not stored in config.json.")
    add_note(doc, "Development fallback only.", "If you are running the Python source or an executable built without the repository secret, set MXCONNECT_API_KEY for that Windows user and restart the program. Do not paste the key into a report, screenshot, config file, or support message.", fill=PALE_BLUE, color=NAVY)

    add_heading(doc, "3. First-time UI setup", 1)
    number = add_bullet_numbering(doc, "decimal", left=540, hanging=270)
    for item in [
        "Double-click TSYS_PAX_BATCH_REPORT.exe.",
        "Confirm Config file points to the config.json beside the program. Use Open config if needed.",
        "Choose the Output folder where CSV files should be written.",
        "Set Report timeframe (days). Use 3 for the normal three-day report.",
        "Click Save config. The log should show the exact path saved.",
    ]:
        add_list_item(doc, item, number, after=4, line=1.25)

    add_heading(doc, "4. Run the report", 1)
    add_body(doc, "Click Run report. The log panel will show authentication, batch and roster retrieval, authorization scanning, record counts, and the output paths. Wait for the completed message and confirm the primary CSV was written before sending or reviewing it.")
    add_body(doc, "The program uses the active MXConnect TSYS roster as its store list. It excludes stores that have an accepted batch in the selected window, then checks approved authorization activity for the remaining candidate stores.")

    add_heading(doc, "5. Understand the output", 1)
    add_table(
        doc,
        ["File", "What to use it for"],
        [
            ("pinpad_batch_not_closed_3_days.csv", "Primary list of stores to review. The number changes if you select a different timeframe."),
            ("in_use_not_batched_detail.csv", "Supporting authorization detail. It can show terminalNumber and approved amounts, but it does not prove that terminal was the one that failed to batch."),
            ("needs_mapping_or_review.csv", "Rows excluded from the primary list because store information was missing or conflicting. Fix the issue before treating them as reportable."),
            ("batch_history.csv", "Accepted batch records used by the current run."),
            ("termid_account_history.csv", "Account/termID history from accepted batch records for reference."),
            ("tsys_batch_report_summary_<N>_days.docx", "One-page summary of the CSV outputs created by the current run."),
        ],
        [2850, 6510],
    )
    add_body(doc, "The primary CSV columns are URL, STORENAME, DEVICE, AMOUNT, and TERMID. TERMID is intentionally blank in this version. AMOUNT is the sum of approved authorization amounts returned for the store account; it is not the batch total.")
    add_body(doc, "Each report run creates a new timestamped subfolder under the configured output folder using YYYYMMDD_HHMMSS. If two runs occur in the same second, the later folder receives a suffix such as _01, so prior outputs are not overwritten.")

    add_heading(doc, "6. Optional display overrides", 1)
    add_body(doc, "The UI can maintain optional rows for URL, Store Name, Device, and Account Number. These values are for display only. Account Number is the TSYS merchant account used to associate a display override with the active roster.")
    add_body(doc, "If there are two enabled rows for the same account with different values, the account is excluded from the primary CSV and placed in needs_mapping_or_review.csv. This prevents a conflicting manual label from reaching the report.")

    add_heading(doc, "7. Scheduled operation", 1)
    add_body(doc, "For Windows Task Scheduler, create a task that runs the CLI executable with the same config file:")
    add_code(doc, '''TSYS_PAX_BATCH_REPORT_CLI.exe --run-report --config "C:\\TSYS_PAX_BATCH_REPORT\\config.json"''')
    add_body(doc, "Set the task's Start in folder to the folder containing the executables. The account running the task needs access to the executables, config.json, and output folder. No API-key environment variable is required when using the official Windows package.")

    add_heading(doc, "8. Historical refresh", 1)
    add_body(doc, "Use Refresh historical data when you want a separate historical snapshot. The default lookback is 90 days. Historical files are written under a timestamped folder inside the output folder's historical subfolder and do not replace the primary report.")

    add_heading(doc, "9. Common problems", 1)
    add_table(
        doc,
        ["Problem", "What to do"],
        [
            ("Set MXCONNECT_API_KEY before running the API report.", "Use the official Windows package, which contains the build-time key. If running source or a local build, set the fallback variable for the same Windows user and restart the program without printing the value."),
            ("Output folder did not change.", "Confirm the Config file path, choose the folder again, and click Save config. Inspect that exact config.json path."),
            ("No rows in the primary CSV.", "This may be correct: no candidate stores had approved authorization activity, or all active stores with activity had an accepted batch."),
            ("A store is in the review CSV.", "Review the reason/details columns. Correct the store name or conflicting account override, then rerun."),
            ("UI says report exited with code 1.", "Read the error immediately above it in the log. Correct the API key, config, network, or API error and rerun."),
        ],
        [2900, 6460],
    )
    add_note(doc, "Safe interpretation.", "Use the primary CSV as a store-level follow-up list. Use the detail CSV to understand the approved authorization exposure. Do not describe a row as a verified unbatched physical device unless separate authoritative evidence is available.", fill=RED_FILL, color=RED)

    add_source_note(doc, "This guide reflects the current Windows build and configuration template in asispirits/GET_TSYS_BATCH.")
    path = OUT_DIR / "TSYS_PAX_BATCH_REPORT_User_Setup_and_Operation_Guide.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    technical = build_technical_manual()
    user = build_user_guide()
    print(technical)
    print(user)
