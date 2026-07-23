#!/usr/bin/env python3
"""Generate the store-level TSYS PAX batch-not-closed report.

The report flags an active TSYS store when it has authorization activity in the
configured absolute window and no batch record for the store's account in the
report window. It links displayed termID values through accountNumber and batch
history, but does not claim that a specific termID or physical device failed to
batch. Optional configuration rows provide display overrides only; MXConnect's
active merchant roster is the authoritative store list.
"""

import csv
import base64
import hashlib
from html import escape
import json
import os
import subprocess
import sys
import tempfile
import threading
import tkinter as tk
import webbrowser
from collections import Counter, defaultdict
from argparse import ArgumentParser
from datetime import datetime, time, timedelta, timezone
from pathlib import Path
from time import perf_counter
from tkinter import filedialog, messagebox, ttk
from urllib.parse import quote

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
import requests

try:
    from embedded_api_key import EMBEDDED_API_KEY
except ImportError:
    EMBEDDED_API_KEY = ""


BASE_URL = "https://api.mxconnect.com"
AUTH_PATH = "/security/v1/apiKey/authenticate"
BATCH_PATH = "/report/v1/tsys/batch/export"
AUTHORIZATION_PATH = "/report/v1/tsys/authorization/export"
UAR_PATH = "/boarding/v1/uar"
TSYS_PRODUCT_ID = "3"
EMAIL_KEYS = [
    "STORENAME",
    "AMOUNT",
    "accountNumber",
    "termID",
    "approvedAmount",
    "approvedCount",
    "authorizationActivityCount",
    "authorizedUnbatchedAmount",
    "batchEvidence",
    "lastBatchDate",
]
REVIEW_KEYS = EMAIL_KEYS + ["reason", "details"]
RAW_BATCH_KEYS = [
    "created", "rejected", "batchNumber", "accountNumber", "termID",
    "salesAmount", "salesCount", "refundAmount", "refundCount", "netAmount",
    "netCount", "PPSNotFundedTotal", "PPSFundedTotal", "bankNumber", "batchDate",
    "fileId", "filePath", "fileName", "id", "accountId", "entityId", "acl",
    "labels", "locationId", "name", "uar", "domain",
]
BATCH_HISTORY_KEYS = [
    "accountNumber", "termID", "batchNumber", "batchDate", "rejected",
    "salesAmount", "salesCount", "refundAmount", "refundCount",
    "netAmount", "netCount",
]
REPORT_CATALOG_FILENAME = "BPOS_REPORT_CATALOG.json"
REPORT_CATALOG_SCHEMA_VERSION = 2
MIN_CONSTANT_REPORTS = 3
TERM_HISTORY_CONSTANT_FIELD = "termHistory"
BATCH_HISTORY_LABEL = "Batch history"
DEFAULT_AUTH_CHECK_DAYS = 3
DEFAULT_LAST_BATCH_LOOKBACK_DAYS = 60


def text(value):
    if value is None:
        return ""
    return str(value).strip()


def parse_amount(value):
    if value is None or value == "":
        return 0.0
    try:
        return float(str(value).replace(",", "").replace("$", "").strip())
    except (TypeError, ValueError):
        return 0.0


def is_approved(record):
    return text(record.get("authorizationResponseStatus")).casefold() == "approved"


def is_accepted_batch(record):
    return text(record.get("rejected")).casefold() in {"no", "false", "0", "n"}


def batch_status(record):
    rejected = text(record.get("rejected")).casefold()
    if not rejected:
        return "unknown"
    if rejected in {"no", "false", "0", "n"}:
        return "accepted"
    if rejected in {"yes", "true", "1", "y"}:
        return "rejected"
    return "unknown"


def clean_api_key(value):
    value = text(value)
    if len(value) >= 2 and value[0] == value[-1] and value[0] in "\"'":
        value = value[1:-1].strip()
    return value


def application_entrypoint():
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve()
    return Path(__file__).resolve()


_PROCESS_OUTPUT_HANDLE = None


def configure_process_output(log_path=None):
    global _PROCESS_OUTPUT_HANDLE
    if log_path:
        path = Path(log_path).expanduser()
        path.parent.mkdir(parents=True, exist_ok=True)
        _PROCESS_OUTPUT_HANDLE = path.open("w", encoding="utf-8", buffering=1)
        sys.stdout = _PROCESS_OUTPUT_HANDLE
        sys.stderr = _PROCESS_OUTPUT_HANDLE
    elif sys.stdout is None or sys.stderr is None:
        _PROCESS_OUTPUT_HANDLE = open(os.devnull, "w", encoding="utf-8")
        sys.stdout = _PROCESS_OUTPUT_HANDLE
        sys.stderr = _PROCESS_OUTPUT_HANDLE


def flush_output():
    if sys.stdout is not None:
        sys.stdout.flush()


def load_config(path):
    with path.open("r", encoding="utf-8") as input_file:
        config = json.load(input_file)
    if not isinstance(config, dict):
        raise ValueError(f"Config must be a JSON object: {path}")
    devices = config.get("devices")
    if not isinstance(devices, list):
        raise ValueError(f"Config must contain a devices array: {path}")
    return config, devices


def resolve_output_directory(config_path, config):
    raw_output = text(config.get("outputDirectory")) or "./tsys-auditdata"
    output_path = Path(raw_output).expanduser()
    if not output_path.is_absolute():
        output_path = config_path.parent / output_path
    return output_path.resolve()


def create_timestamped_output_directory(output_directory):
    output_directory = Path(output_directory)
    timestamp = datetime.now().strftime("%m-%d-%Y %I-%M-%S %p")
    attempt = 0
    while True:
        suffix = f"_{attempt:02d}" if attempt else ""
        run_directory = output_directory / f"{timestamp}{suffix}"
        try:
            run_directory.mkdir(parents=True, exist_ok=False)
            return run_directory
        except FileExistsError:
            attempt += 1


def portable_output_directory(config_path, output_path):
    selected = Path(text(output_path)).expanduser()
    if not selected.is_absolute():
        return text(output_path) or "./tsys-auditdata"
    relative = os.path.relpath(str(selected), str(config_path.parent))
    return relative.replace(os.sep, "/") or "."


def save_config(path, config, devices, output_directory=None):
    updated = dict(config)
    if output_directory is not None:
        updated["outputDirectory"] = portable_output_directory(path, output_directory)
    updated["devices"] = devices
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = path.with_suffix(path.suffix + ".tmp")
    with temporary_path.open("w", encoding="utf-8", newline="\n") as output_file:
        json.dump(updated, output_file, indent=2, ensure_ascii=False)
        output_file.write("\n")
    temporary_path.replace(path)


def write_csv(rows, path, keys):
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = path.with_suffix(path.suffix + ".tmp")
    with temporary_path.open("w", encoding="utf-8", newline="") as output_file:
        writer = csv.DictWriter(output_file, fieldnames=keys, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)
    temporary_path.replace(path)


def write_styled_xlsx(rows, path, keys, left_aligned_columns):
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = path.with_suffix(path.suffix + ".tmp")
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = path.stem[:31]

    for column_index, key in enumerate(keys, start=1):
        cell = worksheet.cell(row=1, column=column_index, value=key)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.font = Font(bold=True)

    for row_index, row in enumerate(rows, start=2):
        for column_index, key in enumerate(keys, start=1):
            cell = worksheet.cell(row=row_index, column=column_index, value=row.get(key, ""))
            cell.alignment = Alignment(
                horizontal="left" if column_index in left_aligned_columns else "center",
                vertical="center",
            )

    widths = [28, 14, 20, 18, 18, 16, 25][:len(keys)]
    for column_index, width in enumerate(widths, start=1):
        worksheet.column_dimensions[chr(64 + column_index)].width = width
    worksheet.freeze_panes = "A2"
    worksheet.auto_filter.ref = worksheet.dimensions
    worksheet.row_dimensions[1].height = 22
    workbook.save(temporary_path)
    temporary_path.replace(path)


def write_primary_xlsx(rows, path):
    write_styled_xlsx(rows, path, EMAIL_KEYS, {7})


class MxConnectClient:
    def __init__(self, base_url, api_key):
        self.base_url = base_url.rstrip("/")
        self.api_key = api_key
        self.session = requests.Session()
        self.session.headers.update({"Accept": "application/json"})
        self.token = None

    def authenticate(self):
        response = self.session.post(
            f"{self.base_url}{AUTH_PATH}",
            json={"value": self.api_key},
            headers={"Content-Type": "application/json"},
            timeout=90,
        )
        if response.status_code != 200:
            body = response.text[:500].replace("\n", " ")
            raise RuntimeError(
                f"MXConnect authentication failed with {response.status_code} "
                f"at {self.base_url}{AUTH_PATH}: {body}"
            )
        data = response.json()
        self.token = text(data.get("token"))
        if not self.token:
            raise RuntimeError("MXConnect authentication response did not contain a token.")

    def request(self, method, url, payload=None, retry=True):
        if not self.token:
            self.authenticate()

        headers = {"Authorization": f"Bearer {self.token}"}
        if method == "POST":
            headers["Content-Type"] = "application/json"
            response = self.session.post(url, json=payload or {}, headers=headers, timeout=120)
        else:
            response = self.session.get(url, headers=headers, timeout=120)

        if response.status_code == 401 and retry:
            self.token = None
            self.authenticate()
            return self.request(method, url, payload, retry=False)

        if response.status_code >= 400:
            body = response.text[:500].replace("\n", " ")
            raise RuntimeError(f"MXConnect request failed with {response.status_code}: {body}")
        return response.json()


def response_page(data):
    if isinstance(data, list):
        return data, None
    if not isinstance(data, dict):
        raise ValueError("Unexpected MXConnect response shape.")
    page = data.get("records")
    if page is None:
        page = data.get("data")
    if page is None:
        page = []
    if not isinstance(page, list):
        raise ValueError("Unexpected MXConnect records response shape.")
    total = data.get("totalRecords")
    try:
        total = int(total) if total is not None else None
    except (TypeError, ValueError):
        total = None
    return page, total


def fetch_all(client, url):
    records = []
    payload = {}
    data = client.request("POST", url, payload)
    expected_total = None

    while True:
        page, total = response_page(data)
        if total is not None:
            expected_total = total
        records.extend(page)
        if not page:
            if expected_total is not None and len(records) != expected_total:
                raise RuntimeError(
                    f"MXConnect returned an incomplete export: received {len(records)} "
                    f"of {expected_total} records."
                )
            break
        if expected_total is not None and len(records) >= expected_total:
            if len(records) != expected_total:
                raise RuntimeError(
                    f"MXConnect returned {len(records)} of {expected_total} records."
                )
            break

        scroll_id = data.get("_scroll_id") or data.get("scrollId")
        if not scroll_id:
            raise RuntimeError("MXConnect returned more records but no scroll ID.")
        payload = {"scrollId": scroll_id}
        data = client.request("POST", url, payload)

    return records


def date_window(days):
    end_time = datetime.combine(datetime.now(timezone.utc).date(), time.min) + timedelta(hours=4)
    return end_time - timedelta(days=days), end_time


def batch_url(base_url, start_date, end_date):
    filter_object = {
        "must": [
            {"bool": {"should": [], "minimum_should_match": 1}},
            {
                "bool": {
                    "should": [
                        {
                            "range": {
                                "batchDate": {
                                    "gte": f"{start_date}T00:00:00.000Z",
                                    "lte": f"{end_date}T23:59:59.999Z",
                                }
                            }
                        }
                    ],
                    "minimum_should_match": 1,
                }
            },
        ]
    }
    encoded_filter = quote(json.dumps(filter_object, separators=(",", ":")), safe="")
    return (
        f"{base_url}{BATCH_PATH}?filter={encoded_filter}&size=100&from=0"
    )


def authorization_url(base_url, auth_window, account_numbers):
    fields = [
        "accountNumber",
        "terminalNumber",
        "authorizedAmount",
        "authorizationResponseStatus",
    ]
    encoded_accounts = quote(json.dumps(sorted(account_numbers)), safe="")
    encoded_fields = quote(json.dumps(fields, separators=(",", ":")), safe="")
    return (
        f"{base_url}{AUTHORIZATION_PATH}?dr_type=q&dr_quick={quote(auth_window)}"
        f"&accountNumbers={encoded_accounts}&fields={encoded_fields}&size=100&from=0"
    )


def authorization_absolute_url(base_url, start_date, end_date, account_numbers, fields=None):
    requested_fields = fields or [
        "accountNumber",
        "terminalNumber",
        "authorizedAmount",
        "authorizationResponseStatus",
    ]
    encoded_accounts = quote(json.dumps(sorted(account_numbers)), safe="")
    encoded_fields = quote(json.dumps(requested_fields, separators=(",", ":")), safe="")
    return (
        f"{base_url}{AUTHORIZATION_PATH}?dr_type=abs"
        f"&dr_from={quote(f'{start_date}T00:00:00.000Z')}"
        f"&dr_to={quote(f'{end_date}T23:59:59.999Z')}"
        f"&accountNumbers={encoded_accounts}&fields={encoded_fields}&size=100&from=0"
    )


def uar_url(base_url, offset):
    return f"{base_url}{UAR_PATH}?size=100&from={offset}"


def fetch_uar(client):
    records = []
    offset = 0
    while True:
        data = client.request("GET", uar_url(client.base_url, offset))
        if isinstance(data, list):
            page = data
            total = None
        else:
            page, total = response_page(data)
        if not page:
            break
        records.extend(page)
        offset += len(page)
        if total is not None and offset >= total:
            break
        if len(page) < 100:
            break
    return records


def active_tsys_accounts(uar_records):
    accounts = set()
    for record in uar_records:
        account = record.get("account") or {}
        if (
            text(account.get("number"))
            and text((account.get("product") or {}).get("id")) == TSYS_PRODUCT_ID
            and account.get("active") is True
            and text(account.get("status")) not in {"Closed", "Terminated", "Suspended"}
        ):
            accounts.add(text(account.get("number")))
    return accounts


def active_tsys_roster(uar_records):
    roster = {}
    for record in uar_records:
        account = record.get("account") or {}
        account_number = text(account.get("number"))
        if not account_number:
            continue
        if (
            text((account.get("product") or {}).get("id")) != TSYS_PRODUCT_ID
            or account.get("active") is not True
            or text(account.get("status")) in {"Closed", "Terminated", "Suspended"}
        ):
            continue
        roster[account_number] = {
            "storeName": text((record.get("location") or {}).get("name")),
            "status": text(account.get("status")),
        }
    return roster


def store_display_overrides(raw_devices):
    """Use configured account-level values only as display overrides.

    Account numbers are optional here. They are not used to validate a device;
    they only allow a manually maintained URL/store-name override to follow a
    known merchant account.
    """
    overrides = {}
    conflicts = set()
    for raw in raw_devices:
        if raw.get("enabled", True) is False:
            continue
        account = text(raw.get("accountNumber"))
        if not account:
            continue
        candidate = {
            "URL": text(raw.get("url")),
            "STORENAME": text(raw.get("storeName")),
            "DEVICE": text(raw.get("device")) or "PAX",
        }
        existing = overrides.get(account)
        if existing is not None and existing != candidate:
            conflicts.add(account)
        else:
            overrides[account] = candidate
    return overrides, conflicts


def create_store_report(
    uar_records,
    raw_devices,
    activity_summary,
    current_batch_records,
    last_batch_by_account=None,
    term_history_by_account=None,
    authorized_unbatched_by_account=None,
    history_floor_date="",
):
    roster = active_tsys_roster(uar_records)
    overrides, override_conflicts = store_display_overrides(raw_devices)
    last_batch_by_account = last_batch_by_account or {}
    term_history_by_account = term_history_by_account or {}
    authorized_unbatched_by_account = authorized_unbatched_by_account or {}
    batched_accounts = {
        text(record.get("accountNumber"))
        for record in current_batch_records
        if text(record.get("accountNumber"))
    }

    email_rows = []
    review_rows = []

    def review_row(reason, details, **values):
        row = {key: "" for key in EMAIL_KEYS}
        row.update(values)
        row["reason"] = reason
        row["details"] = details
        return row

    for account in sorted(activity_summary):
        if account in batched_accounts:
            continue
        term_history = term_history_by_account.get(account, {})
        term_ids = term_history.get("termIDs", [])
        report_last_batch_date = term_history.get("lastBatchDate") or last_batch_by_account.get(
            account, ""
        )
        merchant = roster.get(account)
        if merchant is None:
            review_rows.append(
                review_row(
                    "ACCOUNT_NOT_FOUND_IN_ACTIVE_TSYS_ROSTER",
                    "",
                    accountNumber=account,
                    termID=", ".join(term_ids),
                    lastBatchDate=report_last_batch_date,
                )
            )
            continue

        override = overrides.get(account, {})
        if account in override_conflicts:
            review_rows.append(
                review_row(
                    "CONFLICTING_STORE_DISPLAY_OVERRIDES",
                    "More than one configured display row exists for this account.",
                    STORENAME=merchant["storeName"],
                    accountNumber=account,
                    termID=", ".join(term_ids),
                    lastBatchDate=report_last_batch_date,
                )
            )
            continue

        store_name = override.get("STORENAME") or merchant["storeName"]
        if not store_name:
            review_rows.append(
                review_row(
                    "MISSING_STORE_NAME",
                    "Active TSYS account has no location name.",
                    accountNumber=account,
                    termID=report_term_id,
                    lastBatchDate=report_last_batch_date,
                )
            )
            continue

        summary = activity_summary[account]
        output_terms = term_ids or [None]
        for term_id in output_terms:
            email_rows.append(
                {
                    "STORENAME": store_name,
                    "AMOUNT": f"{authorized_unbatched_by_account.get(account, 0.0):.2f}",
                    "accountNumber": account,
                    "termID": term_id or "NULL",
                    "approvedAmount": f"{summary['approvedAmount']:.2f}",
                    "approvedCount": summary["approvedCount"],
                    "authorizationActivityCount": summary["activityCount"],
                    "authorizedUnbatchedAmount": f"{authorized_unbatched_by_account.get(account, 0.0):.2f}",
                    "batchEvidence": "No batch record in report window",
                    "lastBatchDate": report_last_batch_date or f"No batch in {history_floor_date} lookback",
                }
            )

    return email_rows, review_rows


def build_auth_summary(records, require_approved):
    summary = {}
    for record in records:
        if require_approved and not is_approved(record):
            continue
        account = text(record.get("accountNumber"))
        if not account:
            continue
        terminal = text(record.get("terminalNumber"))
        entry = summary.setdefault(
            (account, terminal), {"amount": 0.0, "count": 0}
        )
        entry["amount"] += parse_amount(record.get("authorizedAmount"))
        entry["count"] += 1
    return summary


def build_account_activity_summary(records):
    """Summarize any authorization activity while retaining approved totals."""
    summary = {}
    for record in records:
        account = text(record.get("accountNumber"))
        if not account:
            continue
        account_summary = summary.setdefault(
            account,
            {
                "activityCount": 0,
                "approvedAmount": 0.0,
                "approvedCount": 0,
                "terminals": set(),
            },
        )
        account_summary["activityCount"] += 1
        terminal = text(record.get("terminalNumber"))
        if terminal:
            account_summary["terminals"].add(terminal)
        if is_approved(record):
            account_summary["approvedAmount"] += parse_amount(record.get("authorizedAmount"))
            account_summary["approvedCount"] += 1
    return summary


def authorized_amount_url(base_url, account_number, start_date, end_date):
    return authorization_absolute_url(
        base_url,
        start_date,
        end_date,
        {account_number},
        fields=["accountNumber", "authorizedAmount", "authorizationResponseStatus"],
    )


def authorized_unbatched_amount(client, base_url, account_number, start_date, end_date):
    records = fetch_all(
        client,
        authorized_amount_url(base_url, account_number, start_date, end_date),
    )
    return sum(
        parse_amount(record.get("authorizedAmount"))
        for record in records
        if is_approved(record)
    )


def iso_date_from_display(value, fallback):
    value = text(value)
    for pattern in ("%m/%d/%Y %I:%M:%S %p", "%Y-%m-%d"):
        try:
            return datetime.strptime(value, pattern).strftime("%Y-%m-%d")
        except ValueError:
            continue
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00")).strftime("%Y-%m-%d")
    except ValueError:
        return fallback


def canonical_json(value):
    return json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    )


def fingerprint_records(records):
    normalized_records = sorted(canonical_json(record) for record in records)
    return hashlib.sha256(
        canonical_json(normalized_records).encode("utf-8")
    ).hexdigest()


class AccountReportCatalog:
    """Persist validated derived report data and account-level constants."""

    def __init__(self, path):
        self.path = Path(path)
        self.valid = True
        self.validation_message = ""
        self._payload = {
            "schemaVersion": REPORT_CATALOG_SCHEMA_VERSION,
            "reportIds": [],
            "accountConstants": {},
            "historyView": None,
        }
        self._load()

    @staticmethod
    def _integrity(payload):
        unsigned = {
            key: value for key, value in payload.items() if key != "integrity"
        }
        return hashlib.sha256(canonical_json(unsigned).encode("utf-8")).hexdigest()

    def _invalidate(self, message):
        self.valid = False
        self.validation_message = message
        self._payload = {
            "schemaVersion": REPORT_CATALOG_SCHEMA_VERSION,
            "reportIds": [],
            "accountConstants": {},
            "historyView": None,
        }

    def _load(self):
        if not self.path.exists():
            return
        try:
            with self.path.open("r", encoding="utf-8") as input_file:
                payload = json.load(input_file)
            self._validate(payload)
            self._payload = {
                key: payload[key]
                for key in ("schemaVersion", "reportIds", "accountConstants", "historyView")
            }
        except (OSError, ValueError, TypeError, KeyError) as error:
            self._invalidate(f"ignored invalid catalog: {error}")

    @staticmethod
    def _validate(payload):
        if not isinstance(payload, dict):
            raise ValueError("catalog root must be an object")
        if payload.get("schemaVersion") != REPORT_CATALOG_SCHEMA_VERSION:
            raise ValueError("unsupported catalog schema")
        if payload.get("integrity") != AccountReportCatalog._integrity(payload):
            raise ValueError("catalog integrity check failed")

        report_ids = payload.get("reportIds")
        if not isinstance(report_ids, list) or any(not isinstance(item, str) for item in report_ids):
            raise ValueError("reportIds must be a list of strings")

        account_constants = payload.get("accountConstants")
        if not isinstance(account_constants, dict):
            raise ValueError("accountConstants must be an object")
        for account, fields in account_constants.items():
            if not isinstance(account, str) or not account or not isinstance(fields, dict):
                raise ValueError("invalid account constant entry")
            for field, tracker in fields.items():
                if not isinstance(field, str) or not isinstance(tracker, dict):
                    raise ValueError("invalid account constant field")
                if tracker.get("value") is None:
                    raise ValueError("constant value cannot be null")
                report_count = tracker.get("reportCount")
                if not isinstance(report_count, int) or report_count < 1:
                    raise ValueError("reportCount must be a positive integer")
                if not isinstance(tracker.get("lastReportId"), str):
                    raise ValueError("lastReportId must be a string")

        history_view = payload.get("historyView")
        if history_view is not None:
            if not isinstance(history_view, dict):
                raise ValueError("historyView must be an object or null")
            if not isinstance(history_view.get("fingerprint"), str):
                raise ValueError("history fingerprint must be a string")
            if not isinstance(history_view.get("reportCount"), int) or history_view["reportCount"] < 1:
                raise ValueError("history reportCount must be a positive integer")
            views = history_view.get("views")
            if not isinstance(views, dict):
                raise ValueError("history views must be an object")
            if not isinstance(views.get("lastBatchByAccount"), dict):
                raise ValueError("lastBatchByAccount must be an object")
            if not isinstance(views.get("termHistoryByAccount"), dict):
                raise ValueError("termHistoryByAccount must be an object")
            if not isinstance(views.get("termHistoryRows"), list):
                raise ValueError("termHistoryRows must be a list")

    def stable_account_values(self, field, accounts=None):
        selected_accounts = set(accounts) if accounts is not None else None
        stable = {}
        for account, fields in self._payload["accountConstants"].items():
            if selected_accounts is not None and account not in selected_accounts:
                continue
            tracker = fields.get(field)
            if tracker and tracker["reportCount"] >= MIN_CONSTANT_REPORTS:
                stable[account] = tracker["value"]
        return stable

    def cached_history_views(self, fingerprint):
        history_view = self._payload.get("historyView")
        if (
            not self.valid
            or not history_view
            or history_view.get("fingerprint") != fingerprint
            or history_view.get("reportCount", 0) < MIN_CONSTANT_REPORTS
        ):
            return None
        return history_view["views"]

    def record_report(self, report_id, account_observations, history_fingerprint, history_views):
        if report_id in self._payload["reportIds"]:
            return False
        self._payload["reportIds"].append(report_id)

        account_constants = self._payload["accountConstants"]
        for account, observations in account_observations.items():
            fields = account_constants.setdefault(account, {})
            for field, value in observations.items():
                if value is None:
                    fields.pop(field, None)
                    continue
                previous = fields.get(field)
                if previous and previous["value"] == value:
                    report_count = previous["reportCount"] + 1
                else:
                    report_count = 1
                fields[field] = {
                    "value": value,
                    "reportCount": report_count,
                    "lastReportId": report_id,
                }

        previous_history = self._payload.get("historyView")
        if previous_history and previous_history["fingerprint"] == history_fingerprint:
            report_count = previous_history["reportCount"] + 1
        else:
            report_count = 1
        self._payload["historyView"] = {
            "fingerprint": history_fingerprint,
            "reportCount": report_count,
            "views": history_views,
        }
        return True

    def save(self):
        self.path.parent.mkdir(parents=True, exist_ok=True)
        payload = dict(self._payload)
        payload["integrity"] = self._integrity(payload)
        temporary_path = self.path.with_suffix(self.path.suffix + ".tmp")
        with temporary_path.open("w", encoding="utf-8", newline="\n") as output_file:
            json.dump(payload, output_file, indent=2, ensure_ascii=False)
            output_file.write("\n")
        temporary_path.replace(self.path)


class BatchHistoryCatalog:
    """Build reusable accepted-batch indexes in one pass."""

    @classmethod
    def from_cached_views(cls, views):
        catalog = cls.__new__(cls)
        catalog.by_pair = {}
        catalog.term_accounts = {}
        catalog.last_batch_by_account = views["lastBatchByAccount"]
        catalog.term_history_by_account = views["termHistoryByAccount"]
        catalog.term_history_rows = views["termHistoryRows"]
        return catalog

    def __init__(self, records):
        latest_by_account = {}
        by_pair = {}
        term_accounts = {}

        for record in records:
            account = text(record.get("accountNumber"))
            if not account:
                continue

            sort_value = text(record.get("batchDate")) or text(record.get("created"))
            formatted_timestamp = None

            existing_account = latest_by_account.get(account)
            if existing_account is None or sort_value > existing_account["_sortValue"]:
                formatted_timestamp = format_batch_timestamp(record)
                latest_by_account[account] = {
                    "lastBatchDate": formatted_timestamp,
                    "_sortValue": sort_value,
                }

            term = text(record.get("termID"))
            if not term:
                continue

            pair = (account, term)
            existing_pair = by_pair.get(pair)
            if existing_pair is None or sort_value > existing_pair["_sortValue"]:
                if formatted_timestamp is None:
                    formatted_timestamp = format_batch_timestamp(record)
                by_pair[pair] = {
                    "lastBatchDate": formatted_timestamp,
                    "_sortValue": sort_value,
                }

            term_accounts.setdefault(term, set()).add(account)

        sorted_pairs = sorted(by_pair.items())
        history = defaultdict(lambda: {
            "termIDs": [],
            "lastBatchDate": "",
            "_sortValue": "",
        })
        term_history_rows = []

        for (account, term), detail in sorted_pairs:
            account_history = history[account]
            account_history["termIDs"].append(term)
            if detail["_sortValue"] > account_history["_sortValue"]:
                account_history["lastBatchDate"] = detail["lastBatchDate"]
                account_history["_sortValue"] = detail["_sortValue"]

            term_history_rows.append(
                {
                    "accountNumber": account,
                    "termID": term,
                    "lastBatchDate": detail["lastBatchDate"],
                }
            )

        self.by_pair = by_pair
        self.term_accounts = term_accounts
        self.last_batch_by_account = {
            account: detail["lastBatchDate"]
            for account, detail in latest_by_account.items()
        }
        self.term_history_by_account = {
            account: {
                "termIDs": detail["termIDs"],
                "lastBatchDate": detail["lastBatchDate"],
            }
            for account, detail in history.items()
        }
        self.term_history_rows = term_history_rows


def batch_index(records):
    catalog = BatchHistoryCatalog(records)
    return catalog.by_pair, catalog.term_accounts


def format_timestamp(raw_timestamp):
    raw_timestamp = text(raw_timestamp)
    if raw_timestamp:
        try:
            parsed = datetime.fromisoformat(raw_timestamp.replace("Z", "+00:00"))
            return parsed.strftime("%m/%d/%Y %I:%M:%S %p")
        except ValueError:
            return raw_timestamp.replace("T", " ").rstrip("Zz")
    return ""


def format_batch_timestamp(record):
    return format_timestamp(text(record.get("batchDate")) or text(record.get("created")))


def compact_batch_history_rows(records):
    rows = []
    for record in records:
        row = {key: text(record.get(key)) for key in BATCH_HISTORY_KEYS}
        row["batchDate"] = format_timestamp(
            text(record.get("batchDate")) or text(record.get("created"))
        )
        rows.append(row)
    return rows


def latest_batch_dates_by_account(records):
    return BatchHistoryCatalog(records).last_batch_by_account


def term_history_by_account(records):
    return BatchHistoryCatalog(records).term_history_by_account


def write_raw_batch_history(records, path):
    keys = list(RAW_BATCH_KEYS)
    for record in records:
        for key in record:
            if key not in keys:
                keys.append(key)
    write_csv(records, path, keys)


def build_term_history_rows(records):
    return BatchHistoryCatalog(records).term_history_rows


def write_term_history(records, path):
    write_csv(
        build_term_history_rows(records),
        path,
        ["accountNumber", "termID", "lastBatchDate"],
    )


def write_term_history_xlsx(records, path):
    by_pair, _ = batch_index(records)
    rows = []
    for (account, term), detail in sorted(by_pair.items()):
        rows.append(
            {
                "accountNumber": account,
                "termID": term,
                "lastBatchDate": detail["lastBatchDate"],
            }
        )
    write_styled_xlsx(rows, path, ["accountNumber", "termID", "lastBatchDate"], {3})


def write_active_roster(records, path):
    rows = []
    for record in records:
        account = record.get("account") or {}
        location = record.get("location") or {}
        product = account.get("product") or {}
        rows.append(
            {
                "accountNumber": text(account.get("number")),
                "storeName": text(location.get("name")),
                "active": account.get("active"),
                "status": text(account.get("status")),
                "productId": text(product.get("id")),
                "processor": text((account.get("processor") or {}).get("name")),
            }
        )
    write_csv(
        rows,
        path,
        ["accountNumber", "storeName", "active", "status", "productId", "processor"],
    )


def read_csv_rows(path):
    if not path:
        return []
    path = Path(path)
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as input_file:
        return list(csv.DictReader(input_file))


def write_summary_docx(path, batch_days, auth_window, report_start, report_end, csv_specs):
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as error:
        raise RuntimeError("python-docx is required to generate the DOCX summary.") from error

    NAVY = "0B2545"
    BLUE = "2E74B5"
    DARK_BLUE = "1F4D78"
    INK = "1F2328"
    MUTED = "5B6573"
    LIGHT_BLUE = "E8EEF5"
    LIGHT_GRAY = "F6F8FA"
    BORDER = "D0D7DE"
    CAUTION_FILL = "FFF7D6"
    CAUTION = "7A5A00"
    RED_FILL = "FDECEC"
    RED = "9B1C1C"
    CONTENT_WIDTH_DXA = 9360
    TABLE_INDENT_DXA = 120

    def set_font(run, name="Calibri", size=10.5, color=INK, bold=None, italic=None):
        run.font.name = name
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic

    def set_cell_shading(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
            node = tc_mar.find(qn(f"w:{name}"))
            if node is None:
                node = OxmlElement(f"w:{name}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def set_cell_width(cell, width):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(width))
        tc_w.set(qn("w:type"), "dxa")

    def set_table_geometry(table, widths):
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(sum(widths)))
        tbl_w.set(qn("w:type"), "dxa")
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
        tbl_ind.set(qn("w:type"), "dxa")
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            node = borders.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                borders.append(node)
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "4")
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), BORDER)
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row_index, row in enumerate(table.rows):
            if row_index == 0:
                tr_pr = row._tr.get_or_add_trPr()
                header = OxmlElement("w:tblHeader")
                header.set(qn("w:val"), "true")
                tr_pr.append(header)
            for cell, width in zip(row.cells, widths):
                set_cell_width(cell, width)
                set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if row_index == 0:
                    set_cell_shading(cell, LIGHT_BLUE)

    def add_para(doc, value, size=10.5, color=INK, bold=False, italic=False, after=6, align=None):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.line_spacing = 1.25
        if align is not None:
            paragraph.alignment = align
        run = paragraph.add_run(str(value))
        set_font(run, size=size, color=color, bold=bold, italic=italic)
        return paragraph

    def add_heading(doc, value, level=1):
        paragraph = doc.add_paragraph(value, style=f"Heading {level}")
        return paragraph

    def add_note(doc, label, value, fill=LIGHT_GRAY, label_color=NAVY):
        table = doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [CONTENT_WIDTH_DXA])
        cell = table.cell(0, 0)
        set_cell_shading(cell, fill)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        label_run = paragraph.add_run(label + " ")
        set_font(label_run, size=10.5, color=label_color, bold=True)
        value_run = paragraph.add_run(value)
        set_font(value_run, size=10.5, color=INK)
        add_para(doc, "", after=2)

    def add_table(doc, headers, rows, widths):
        table = doc.add_table(rows=1, cols=len(headers))
        for cell, header in zip(table.rows[0].cells, headers):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(header)
            set_font(run, size=9.2, color=NAVY, bold=True)
        for row_data in rows:
            cells = table.add_row().cells
            for cell, value in zip(cells, row_data):
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                run = paragraph.add_run(str(value))
                set_font(run, size=9.2, color=INK)
        set_table_geometry(table, widths)
        add_para(doc, "", after=2)
        return table

    def add_page_field(paragraph):
        run = paragraph.add_run("Page ")
        set_font(run, size=8.5, color=MUTED)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instruction)
        run._r.append(end)

    csv_data = []
    for label, csv_path, description in csv_specs:
        if csv_path:
            path_value = Path(csv_path)
            csv_data.append((label, path_value, description, read_csv_rows(path_value)))
    data_by_label = {label: rows for label, _, _, rows in csv_data}
    primary_rows = data_by_label.get("Primary exception report", [])
    review_rows = data_by_label.get("Store review", [])
    detail_rows = data_by_label.get("Authorization detail", []) or [
        row for row in primary_rows
        if text(row.get("accountNumber")) or text(row.get("termID")) or text(row.get("terminalNumber"))
    ]
    batch_rows = data_by_label.get(BATCH_HISTORY_LABEL, [])
    raw_batch_rows = data_by_label.get("Raw batch export", [])
    term_rows = data_by_label.get("Term/account history", [])

    seen_primary_accounts = set()
    unique_primary_rows = []
    for row in primary_rows:
        account = text(row.get("accountNumber")) or "|".join(
            [text(row.get("STORENAME")), text(row.get("DEVICE"))]
        )
        if account in seen_primary_accounts:
            continue
        seen_primary_accounts.add(account)
        unique_primary_rows.append(row)
    primary_total = sum(parse_amount(row.get("AMOUNT")) for row in unique_primary_rows)
    detail_total = sum(parse_amount(row.get("approvedAmount")) for row in detail_rows)
    reason_counts = Counter(text(row.get("reason")) or "UNSPECIFIED" for row in review_rows)
    term_accounts = defaultdict(set)
    for row in batch_rows:
        term = text(row.get("termID"))
        account = text(row.get("accountNumber"))
        if term and account:
            term_accounts[term].add(account)
    ambiguous_terms = [
        (term, ", ".join(sorted(accounts)))
        for term, accounts in sorted(term_accounts.items())
        if len(accounts) > 1
    ]
    rejected_raw = sum(1 for row in raw_batch_rows if not is_accepted_batch(row)) if raw_batch_rows else 0
    batch_sales_total = sum(parse_amount(row.get("salesAmount")) for row in batch_rows)
    batch_net_total = sum(parse_amount(row.get("netAmount")) for row in batch_rows)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05
    for style_name, size, color, before, after in (
        ("Heading 1", 11.5, BLUE, 6, 3),
        ("Heading 2", 10.5, BLUE, 5, 2),
        ("Heading 3", 10, DARK_BLUE, 4, 2),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(3)
    header_run = header.add_run("TSYS_PAX_BATCH_REPORT | Run Summary")
    set_font(header_run, size=7.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Bottle POS | ")
    set_font(footer_run, size=7.5, color=MUTED)
    add_page_field(footer)

    logo_root = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
    logo_path = logo_root / "bottlepos_logo.png"
    if logo_path.exists():
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo_paragraph.paragraph_format.space_after = Pt(1)
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(str(logo_path), width=Inches(1.35))

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    title_run = title.add_run("TSYS_PAX_BATCH_REPORT")
    set_font(title_run, size=18, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(4)
    subtitle_run = subtitle.add_run("Run summary")
    set_font(subtitle_run, size=10.5, color=MUTED)
    add_para(
        doc,
        f"Generated {datetime.now():%Y-%m-%d %H:%M} | Batch window {report_start:%Y-%m-%d} to {report_end:%Y-%m-%d} | Authorization {auth_window}",
        size=8.5,
        color=MUTED,
        after=5,
    )

    if primary_rows:
        add_note(
            doc,
            "Primary finding.",
            f"{len(primary_rows)} row(s) met the report criteria, representing ${primary_total:,.2f} in approved authorized-unbatched exposure. The primary report is store/account level; it does not prove which physical terminal failed to batch.",
            fill=CAUTION_FILL,
            label_color=CAUTION,
        )
    else:
        add_note(
            doc,
            "Primary finding.",
            "No stores met the report criteria for this run. No active TSYS account had authorization activity and no batch record in the selected batch window.",
            fill=LIGHT_GRAY,
            label_color=NAVY,
        )

    add_heading(doc, "Key results", 1)
    add_table(
        doc,
        ["Measure", "Result"],
        [
            ("Stores flagged", f"{len(primary_rows):,}"),
            ("Authorized-unbatched exposure", f"${primary_total:,.2f}"),
            ("Authorization detail", f"{len(detail_rows):,} row(s) / ${detail_total:,.2f}"),
            ("Review rows excluded", f"{len(review_rows):,}"),
            ("Batch records used", f"{len(batch_rows):,}"),
            ("Rejected raw batches", f"{rejected_raw:,}" if raw_batch_rows else "Not requested"),
            ("Reused termIDs", f"{len(ambiguous_terms):,}"),
        ],
        [5200, 4160],
    )

    add_heading(doc, "Important interpretation", 1)
    interpretation = "The primary report is account/term level. termID values are linked through accountNumber to batch history; lastBatchDate is the latest batch timestamp in that history. Authorized-unbatched amounts are account-level approved exposure and are not proof that a specific physical terminal failed to batch."
    if review_rows:
        reason_text = ", ".join(f"{reason}: {count}" for reason, count in reason_counts.most_common(4))
        interpretation += f" Review exclusions by reason: {reason_text}."
    if ambiguous_terms:
        interpretation += f" {len(ambiguous_terms):,} termID value(s) appeared under multiple account numbers in batch history."
    add_para(doc, interpretation, size=8.8, after=4)

    add_heading(doc, "Top flagged stores", 1)
    if primary_rows:
        top_primary = sorted(unique_primary_rows, key=lambda row: parse_amount(row.get("AMOUNT")), reverse=True)[:5]
        add_table(
            doc,
            ["Store", "Account", "Last batch", "Approved amount"],
            [
                (
                    text(row.get("STORENAME")) or "(unnamed store)",
                    text(row.get("accountNumber")) or "-",
                    text(row.get("lastBatchDate")) or "No history",
                    f"${parse_amount(row.get('AMOUNT')):,.2f}",
                )
                for row in top_primary
            ],
            [3900, 2300, 1900, 1660],
        )
        if len(unique_primary_rows) > len(top_primary):
            add_para(doc, f"Showing the top {len(top_primary)} by approved amount; the complete list is in the primary CSV.", size=8.2, color=MUTED, italic=True, after=3)
    else:
        add_para(doc, "No primary exception rows were generated.", size=8.8, after=4)

    add_heading(doc, "Output files", 1)
    file_list = "; ".join(
        f"{path_value.name} ({len(rows):,} rows)"
        for _, path_value, _, rows in csv_data
    )
    add_para(
        doc,
        f"All outputs are saved together in the timestamped run folder: {file_list}. The CSV files remain the detailed source artifacts for audit and follow-up.",
        size=7.8,
        color=MUTED,
        italic=True,
        after=0,
    )

    properties = doc.core_properties
    properties.title = "TSYS_PAX_BATCH_REPORT Summary"
    properties.subject = "Summary of current TSYS/PAX batch report CSV outputs"
    properties.author = "ASI Spirits"
    properties.keywords = "TSYS, PAX, batch report, CSV summary, operations"
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def write_summary_html(path, batch_days, auth_window, report_start, report_end, csv_specs):
    csv_data = []
    for label, csv_path, description in csv_specs:
        if csv_path:
            path_value = Path(csv_path)
            csv_data.append((label, path_value, description, read_csv_rows(path_value)))
    data_by_label = {label: rows for label, _, _, rows in csv_data}
    primary_rows = data_by_label.get("Primary exception report", [])
    review_rows = data_by_label.get("Store review", [])
    detail_rows = data_by_label.get("Authorization detail", []) or [
        row for row in primary_rows
        if text(row.get("accountNumber")) or text(row.get("termID")) or text(row.get("terminalNumber"))
    ]
    batch_rows = data_by_label.get(BATCH_HISTORY_LABEL, [])
    raw_batch_rows = data_by_label.get("Raw batch export", [])

    unique_primary_rows = []
    seen_accounts = set()
    for row in primary_rows:
        key = text(row.get("accountNumber")) or "|".join(
            [text(row.get("STORENAME")), text(row.get("DEVICE"))]
        )
        if key in seen_accounts:
            continue
        seen_accounts.add(key)
        unique_primary_rows.append(row)

    primary_total = sum(parse_amount(row.get("AMOUNT")) for row in unique_primary_rows)
    detail_total = sum(parse_amount(row.get("approvedAmount")) for row in detail_rows)
    reason_counts = Counter(text(row.get("reason")) or "UNSPECIFIED" for row in review_rows)
    term_accounts = defaultdict(set)
    for row in batch_rows:
        term = text(row.get("termID"))
        account = text(row.get("accountNumber"))
        if term and account:
            term_accounts[term].add(account)
    ambiguous_terms = [
        (term, ", ".join(sorted(accounts)))
        for term, accounts in sorted(term_accounts.items())
        if len(accounts) > 1
    ]
    rejected_raw = sum(1 for row in raw_batch_rows if not is_accepted_batch(row)) if raw_batch_rows else 0

    logo_root = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
    logo_path = logo_root / "bottlepos_logo.png"
    logo_html = ""
    if logo_path.exists():
        logo_data = base64.b64encode(logo_path.read_bytes()).decode("ascii")
        logo_html = f'<img class="logo" src="data:image/png;base64,{logo_data}" alt="Bottle POS">'

    def cell(value):
        return escape(text(value) or "-")

    finding = (
        f"{len(unique_primary_rows):,} row(s) met the report criteria, representing "
        f"${primary_total:,.2f} in authorized-unbatched exposure. The report is account/term level; "
        "it does not prove which physical terminal failed to batch."
        if unique_primary_rows
        else "No accounts met the report criteria for this run. No active TSYS account had authorization activity and no batch record in the selected window."
    )
    reason_text = ", ".join(
        f"{escape(reason)}: {count:,}" for reason, count in reason_counts.most_common(4)
    ) or "None"
    ambiguous_text = (
        f"{len(ambiguous_terms):,} termID value(s) appeared under multiple account numbers in batch history."
        if ambiguous_terms else "No reused termID values were found in the batch history."
    )
    top_rows = sorted(
        unique_primary_rows,
        key=lambda row: parse_amount(row.get("AMOUNT") or row.get("approvedAmount")),
        reverse=True,
    )[:5]
    top_rows_html = "".join(
        "<tr>"
        f"<td>{cell(row.get('STORENAME'))}</td>"
        f"<td>{cell(row.get('accountNumber'))}</td>"
        f"<td>{cell(row.get('lastBatchDate'))}</td>"
        f"<td class=amount>${parse_amount(row.get('AMOUNT') or row.get('approvedAmount')):,.2f}</td>"
        "</tr>"
        for row in top_rows
    ) or '<tr><td colspan="4" class="empty">No primary exception rows were generated.</td></tr>'
    file_list_html = "".join(
        f'<li><span>{escape(path_value.name)}</span><small>{len(rows):,} row(s)</small></li>'
        for _, path_value, _, rows in csv_data
    )
    generated = datetime.now().strftime("%Y-%m-%d %H:%M")

    html_document = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>BottlePOS PAX Batch Report</title>
  <style>
    :root {{
      --navy: #203239;
      --navy-2: #17272d;
      --blue: #3d8fca;
      --blue-light: #d9edf7;
      --green: #8bbd35;
      --yellow: #f1b735;
      --red: #df6268;
      --ink: #30383d;
      --muted: #7a8790;
      --border: #d9dde0;
      --surface: #ffffff;
      --canvas: #f4f5f6;
    }}
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; background: var(--canvas); color: var(--ink); font-family: Arial, Helvetica, sans-serif; }}
    .topbar {{ position: relative; height: 72px; display: flex; align-items: center; gap: 18px; padding: 0 30px; background: var(--surface); border-bottom: 1px solid var(--border); }}
    .logo {{ width: 170px; height: auto; display: block; }}
    .brand {{ position: absolute; left: 50%; transform: translateX(-50%); font-size: 20px; font-weight: 700; color: var(--navy); letter-spacing: .2px; white-space: nowrap; }}
    .run-meta {{ margin-left: auto; color: var(--muted); font-size: 12px; text-align: right; line-height: 1.5; }}
    .layout {{ min-height: calc(100vh - 72px); }}
    .main {{ width: 100%; max-width: 1280px; margin: 0 auto; padding: 26px 32px 38px; }}
    .page-title {{ display: flex; align-items: baseline; gap: 12px; margin-bottom: 4px; }}
    h1 {{ margin: 0; font-size: 27px; color: var(--navy); }}
    h2 {{ margin: 26px 0 11px; font-size: 18px; color: var(--navy); }}
    .subtitle {{ color: var(--muted); font-size: 13px; }}
    .finding {{ margin-top: 22px; padding: 17px 20px; background: #fff8df; border: 1px solid #efdaa0; border-left: 5px solid var(--yellow); line-height: 1.5; }}
    .finding strong {{ color: #9a6b00; }}
    .cards {{ display: grid; grid-template-columns: repeat(5, minmax(0, 1fr)); gap: 12px; }}
    .card {{ background: var(--surface); border: 1px solid var(--border); padding: 15px 16px; min-height: 92px; }}
    .card .label {{ color: var(--muted); font-size: 12px; text-transform: uppercase; letter-spacing: .4px; }}
    .card .value {{ margin-top: 9px; color: var(--blue); font-size: 23px; font-weight: 700; }}
    .card.green .value {{ color: var(--green); }}
    .card.yellow .value {{ color: var(--yellow); }}
    .card.red .value {{ color: var(--red); }}
    .panel {{ background: var(--surface); border: 1px solid var(--border); }}
    .panel-title {{ padding: 11px 14px; background: linear-gradient(#fff, #f1f1f1); border-bottom: 1px solid var(--border); font-weight: 700; color: var(--navy); }}
    table {{ width: 100%; border-collapse: collapse; font-size: 13px; }}
    th {{ padding: 10px 12px; text-align: left; background: var(--blue-light); color: var(--navy); border-bottom: 1px solid #bdd7e7; }}
    td {{ padding: 9px 12px; border-bottom: 1px solid #e7e9ea; }}
    tr:last-child td {{ border-bottom: 0; }}
    .amount {{ text-align: right; font-weight: 700; color: var(--green); }}
    .empty {{ color: var(--muted); text-align: center; padding: 18px; }}
    .two-col {{ display: grid; grid-template-columns: minmax(0, 1.35fr) minmax(280px, .65fr); gap: 18px; align-items: start; }}
    .note {{ padding: 14px 16px; background: #f8fafb; border: 1px solid var(--border); color: var(--muted); line-height: 1.5; font-size: 13px; }}
    .files {{ list-style: none; margin: 0; padding: 0; }}
    .files li {{ display: flex; justify-content: space-between; gap: 12px; padding: 10px 14px; border-bottom: 1px solid #e7e9ea; font-size: 12px; }}
    .files li:last-child {{ border-bottom: 0; }}
    .files small {{ color: var(--muted); white-space: nowrap; }}
    footer {{ margin-top: 28px; color: var(--muted); font-size: 12px; text-align: right; }}
    @media (max-width: 900px) {{ .brand {{ position: static; transform: none; flex: 1; text-align: center; }} .main {{ padding: 20px; }} .cards {{ grid-template-columns: repeat(2, minmax(0, 1fr)); }} .two-col {{ grid-template-columns: 1fr; }} }}
    @media print {{ body {{ background: #fff; }} .main {{ max-width: none; padding: 0; }} .topbar {{ padding: 0 0 14px; }} .panel, .card {{ break-inside: avoid; }} }}
  </style>
</head>
<body>
  <header class="topbar">{logo_html}<div class="brand">PAX BATCH REPORT</div><div class="run-meta">Run summary<br>Generated {escape(generated)}</div></header>
  <div class="layout">
    <main class="main">
      <div class="page-title"><h1>Batch exception summary</h1><span class="subtitle">Bottle POS</span></div>
      <div class="subtitle">Batch window: {escape(report_start.strftime('%Y-%m-%d'))} to {escape(report_end.strftime('%Y-%m-%d'))} &nbsp;|&nbsp; Authorization: {escape(auth_window)}</div>
      <section class="finding"><strong>Primary finding.</strong> {escape(finding)}</section>
      <h2>Key results</h2>
      <section class="cards">
        <div class="card"><div class="label">Stores flagged</div><div class="value">{len(unique_primary_rows):,}</div></div>
        <div class="card green"><div class="label">Authorized-unbatched</div><div class="value">${primary_total:,.2f}</div></div>
        <div class="card"><div class="label">Authorization detail</div><div class="value">{len(detail_rows):,}</div></div>
        <div class="card yellow"><div class="label">Batch records</div><div class="value">{len(batch_rows):,}</div></div>
        <div class="card red"><div class="label">Review excluded</div><div class="value">{len(review_rows):,}</div></div>
      </section>
      <div class="two-col">
        <section><h2>Top flagged accounts</h2><div class="panel"><div class="panel-title">Authorized-unbatched exposure</div><table><thead><tr><th>Store</th><th>Account</th><th>Last batch</th><th class=amount>Amount</th></tr></thead><tbody>{top_rows_html}</tbody></table></div></section>
        <section><h2>Output files</h2><div class="panel"><ul class="files">{file_list_html}</ul></div></section>
      </div>
      <h2>Important interpretation</h2>
      <div class="note">The primary report is account/term level. termID values are linked through accountNumber to batch history; lastBatchDate is the latest batch timestamp in that history. Authorized-unbatched amounts are account-level approved exposure, not proof that a specific physical terminal failed to batch. Review exclusions by reason: {reason_text}. {escape(ambiguous_text)}</div>
      <footer>Bottle POS</footer>
    </main>
  </div>
</body>
</html>
"""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(html_document, encoding="utf-8")
    return path


def write_interactive_summary_html(
    path,
    batch_days,
    auth_window,
    report_start,
    report_end,
    data_specs,
    report_context=None,
):
    data_by_label = {label: rows for label, _, _, rows in data_specs}
    filename_by_label = {label: filename for label, filename, _, _ in data_specs}
    primary_rows = data_by_label.get("Primary exception report", [])
    review_rows = data_by_label.get("Store review", [])
    batch_rows = data_by_label.get(BATCH_HISTORY_LABEL, [])

    unique_primary_rows = []
    seen_accounts = set()
    for row in primary_rows:
        key = text(row.get("accountNumber")) or "|".join(
            [text(row.get("STORENAME")), text(row.get("DEVICE"))]
        )
        if key in seen_accounts:
            continue
        seen_accounts.add(key)
        unique_primary_rows.append(row)

    detail_rows = [
        row for row in primary_rows
        if text(row.get("accountNumber")) or text(row.get("termID")) or text(row.get("terminalNumber"))
    ]
    primary_total = sum(parse_amount(row.get("AMOUNT")) for row in unique_primary_rows)
    detail_total = sum(parse_amount(row.get("approvedAmount")) for row in detail_rows)
    reason_counts = Counter(text(row.get("reason")) or "UNSPECIFIED" for row in review_rows)
    term_accounts = defaultdict(set)
    for row in batch_rows:
        term = text(row.get("termID"))
        account = text(row.get("accountNumber"))
        if term and account:
            term_accounts[term].add(account)
    ambiguous_terms = [
        (term, ", ".join(sorted(accounts)))
        for term, accounts in sorted(term_accounts.items())
        if len(accounts) > 1
    ]
    batch_status_counts = Counter(batch_status(row) for row in batch_rows)

    def columns_for(rows, fallback):
        return list(rows[0].keys()) if rows else list(fallback)

    datasets = {
        "primary": {
            "label": "PINPAD BATCH NOT CLOSED",
            "filename": filename_by_label.get("Primary exception report", "PINPAD_BATCH_NOT_CLOSED"),
            "columns": columns_for(primary_rows, EMAIL_KEYS),
            "rows": primary_rows,
        },
        "review": {
            "label": "NEEDS MAPPING OR REVIEW",
            "filename": filename_by_label.get("Store review", "NEEDS_MAPPING_OR_REVIEW"),
            "columns": columns_for(review_rows, REVIEW_KEYS),
            "rows": review_rows,
        },
        "term_history": {
            "label": "TERMID ACCOUNT HISTORY",
            "filename": filename_by_label.get("Term/account history", "TERMID_ACCOUNT_HISTORY"),
            "columns": columns_for( data_by_label.get("Term/account history", []), ["accountNumber", "termID", "lastBatchDate"]),
            "rows": data_by_label.get("Term/account history", []),
        },
        "batch_history": {
            "label": "BATCH HISTORY",
            "filename": filename_by_label.get(BATCH_HISTORY_LABEL, "BATCH_HISTORY"),
            "columns": columns_for(batch_rows, BATCH_HISTORY_KEYS),
            "rows": batch_rows,
        },
    }

    finding = (
        f"{len(unique_primary_rows):,} row(s) met the report criteria, representing "
        f"${primary_total:,.2f} in authorized-unbatched exposure. The report is account/term level; "
        "it does not prove which physical terminal failed to batch."
        if unique_primary_rows
        else "No accounts met the report criteria for this run. No active TSYS account had authorization activity and no batch record in the selected window."
    )
    reason_text = ", ".join(
        f"{reason}: {count:,}" for reason, count in reason_counts.most_common(4)
    ) or "None"
    ambiguous_text = (
        f"{len(ambiguous_terms):,} termID value(s) appeared under multiple account numbers in batch history."
        if ambiguous_terms else "No reused termID values were found in the batch history."
    )
    generated = datetime.now().strftime("%Y-%m-%d %H:%M")

    logo_root = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
    logo_path = logo_root / "bottlepos_logo.png"
    logo_html = ""
    if logo_path.exists():
        logo_data = base64.b64encode(logo_path.read_bytes()).decode("ascii")
        logo_html = f'<img class="logo" src="data:image/png;base64,{logo_data}" alt="Bottle POS">'

    report_data_json = json.dumps(datasets, ensure_ascii=False).replace("<", "\\u003c").replace(">", "\\u003e").replace("&", "\\u0026")
    report_meta_json = json.dumps(
        {
            "generated": generated,
            "batchStart": report_start.strftime("%Y-%m-%d"),
            "batchEnd": report_end.strftime("%Y-%m-%d"),
            "authWindow": auth_window,
            "finding": finding,
            "primaryStoreCount": len(unique_primary_rows),
            "primaryRowCount": len(primary_rows),
            "primaryTotal": primary_total,
            "detailRowCount": len(detail_rows),
            "detailTotal": detail_total,
            "batchRecordCount": len(batch_rows),
            "batchStatusCounts": dict(batch_status_counts),
            "reviewCount": len(review_rows),
            "reasonText": reason_text,
            "ambiguousText": ambiguous_text,
            "ruleSummary": report_context or {},
        },
        ensure_ascii=False,
    ).replace("<", "\\u003c").replace(">", "\\u003e").replace("&", "\\u0026")

    html_document = """<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>BottlePOS PAX Batch Report</title>
  <style>
    :root {
      --navy: #203239;
      --blue: #3d8fca;
      --blue-light: #d9edf7;
      --green: #8bbd35;
      --yellow: #f1b735;
      --red: #df6268;
      --ink: #30383d;
      --muted: #7a8790;
      --border: #d9dde0;
      --surface: #ffffff;
      --canvas: #f4f5f6;
    }
    * { box-sizing: border-box; }
    body { margin: 0; background: var(--canvas); color: var(--ink); font-family: Arial, Helvetica, sans-serif; }
    .topbar { position: relative; height: 72px; display: flex; align-items: center; gap: 18px; padding: 0 30px; background: var(--surface); border-bottom: 1px solid var(--border); }
    .logo { width: 170px; height: auto; display: block; }
    .brand { position: absolute; left: 50%; transform: translateX(-50%); font-size: 20px; font-weight: 700; color: var(--navy); letter-spacing: .2px; white-space: nowrap; }
    .run-meta { margin-left: auto; color: var(--muted); font-size: 12px; text-align: right; line-height: 1.5; }
    .main { width: 100%; max-width: 1440px; margin: 0 auto; padding: 26px 32px 38px; }
    .page-title { display: flex; align-items: baseline; gap: 12px; margin-bottom: 4px; }
    h1 { margin: 0; font-size: 27px; color: var(--navy); }
    h2 { margin: 26px 0 11px; font-size: 18px; color: var(--navy); }
    .subtitle { color: var(--muted); font-size: 13px; }
    .tabs { display: flex; flex-wrap: wrap; gap: 7px; margin-top: 22px; border-bottom: 1px solid var(--border); }
    .tab { border: 1px solid var(--border); border-bottom: 0; border-radius: 5px 5px 0 0; padding: 10px 14px; background: #e9eef1; color: var(--navy); font-weight: 700; cursor: pointer; }
    .tab:hover { background: #d9edf7; }
    .tab.active { background: var(--blue); color: #fff; border-color: var(--blue); }
    .panel { background: var(--surface); border: 1px solid var(--border); }
    .panel-title { padding: 11px 14px; background: linear-gradient(#fff, #f1f1f1); border-bottom: 1px solid var(--border); font-weight: 700; color: var(--navy); }
    .finding { margin-top: 22px; padding: 17px 20px; background: #fff8df; border: 1px solid #efdaa0; border-left: 5px solid var(--yellow); line-height: 1.5; }
    .finding strong { color: #9a6b00; }
    .cards { display: grid; grid-template-columns: repeat(5, minmax(0, 1fr)); gap: 12px; }
    .card { background: var(--surface); border: 1px solid var(--border); padding: 15px 16px; min-height: 92px; }
    .card .label { color: var(--muted); font-size: 12px; text-transform: uppercase; letter-spacing: .4px; }
    .card .value { margin-top: 9px; color: var(--blue); font-size: 23px; font-weight: 700; }
    .card.green .value { color: var(--green); }
    .card.yellow .value { color: var(--yellow); }
    .card.red .value { color: var(--red); }
    .two-col { display: grid; grid-template-columns: minmax(0, 1.35fr) minmax(280px, .65fr); gap: 18px; align-items: start; }
    .note { padding: 14px 16px; background: #f8fafb; border: 1px solid var(--border); color: var(--muted); line-height: 1.5; font-size: 13px; }
    .toolbar { display: flex; align-items: center; gap: 12px; margin: 18px 0 10px; }
    .search { flex: 1; min-width: 220px; padding: 10px 12px; border: 1px solid var(--border); border-radius: 4px; font-size: 14px; }
    .export-button { padding: 10px 14px; border: 1px solid var(--blue); border-radius: 4px; background: var(--blue); color: #fff; font-weight: 700; cursor: pointer; white-space: nowrap; }
    .export-button:hover { background: #2f78aa; }
    .row-count { color: var(--muted); font-size: 13px; white-space: nowrap; }
    .table-wrap { max-height: calc(100vh - 260px); overflow: auto; background: var(--surface); border: 1px solid var(--border); }
    table { width: 100%; border-collapse: collapse; font-size: 13px; }
    th { position: sticky; top: 0; z-index: 1; padding: 10px 12px; text-align: left; background: var(--blue-light); color: var(--navy); border-bottom: 1px solid #bdd7e7; cursor: pointer; white-space: nowrap; }
    th:hover { background: #c5e3f0; }
    td { padding: 9px 12px; border-bottom: 1px solid #e7e9ea; white-space: nowrap; }
    tr:last-child td { border-bottom: 0; }
    tr:hover td { background: #f8fafb; }
    .amount { text-align: right; font-weight: 700; color: var(--green); }
    .empty { color: var(--muted); text-align: center; padding: 18px; }
    footer { margin-top: 28px; color: var(--muted); font-size: 12px; text-align: right; }
    @media (max-width: 900px) { .brand { position: static; transform: none; flex: 1; text-align: center; } .main { padding: 20px; } .cards { grid-template-columns: repeat(2, minmax(0, 1fr)); } .two-col { grid-template-columns: 1fr; } }
    @media print { body { background: #fff; } .main { max-width: none; padding: 0; } .topbar { padding: 0 0 14px; } .tabs, .toolbar { display: none; } .table-wrap { max-height: none; overflow: visible; } .panel, .card { break-inside: avoid; } }
  </style>
</head>
<body>
  <header class="topbar">__LOGO__<div class="brand">PAX BATCH REPORT</div><div class="run-meta">Interactive report<br>Generated __GENERATED__</div></header>
  <main class="main">
    <div class="page-title"><h1>BottlePOS PAX Batch Report</h1><span class="subtitle">Bottle POS</span></div>
    <div class="subtitle">Batch window: __BATCH_START__ to __BATCH_END__ &nbsp;|&nbsp; Authorization: __AUTH_WINDOW__</div>
    <nav id="tabs" class="tabs" role="tablist" aria-label="Report sections"></nav>
    <div id="content"></div>
    <footer>Bottle POS</footer>
  </main>
  <script>
    const REPORT_DATA = __REPORT_DATA__;
    const REPORT_META = __REPORT_META__;
    const TAB_DEFINITIONS = [
      { id: "summary", label: "Summary" },
      { id: "primary", label: "PINPAD BATCH NOT CLOSED" },
      { id: "review", label: "NEEDS MAPPING OR REVIEW" },
      { id: "term_history", label: "TERMID ACCOUNT HISTORY" },
      { id: "batch_history", label: "BATCH HISTORY" },
    ];
    if (REPORT_DATA.raw_batch) TAB_DEFINITIONS.push({ id: "raw_batch", label: "RAW BATCH EXPORT" });

    const tabsElement = document.getElementById("tabs");
    const contentElement = document.getElementById("content");
    const state = {};

    function escapeHtml(value) {
      return String(value ?? "").replace(/[&<>\"]/g, character => ({ "&": "&amp;", "<": "&lt;", ">": "&gt;", '"': "&quot;" }[character]));
    }

    function formatNumber(value) {
      return Number(value || 0).toLocaleString(undefined, { maximumFractionDigits: 2 });
    }

    function formatMoney(value) {
      return "$" + Number(value || 0).toLocaleString(undefined, { minimumFractionDigits: 2, maximumFractionDigits: 2 });
    }

    function csvCell(value) {
      return `"${String(value ?? "").replace(/"/g, '""')}"`;
    }

    function exportFilename(dataset) {
      const baseName = String(dataset.filename || dataset.label || "dataset").split(/[\\/]/).pop();
      return /\\.csv$/i.test(baseName) ? baseName : `${baseName}.csv`;
    }

    function getVisibleRows(datasetId) {
      const dataset = REPORT_DATA[datasetId];
      const currentState = state[datasetId];
      const search = currentState.search.toLowerCase();
      const rows = dataset.rows.filter(row => !search || dataset.columns.some(column => String(row[column] ?? "").toLowerCase().includes(search)));
      if (currentState.sortColumn) {
        rows.sort((left, right) => {
          const leftValue = String(left[currentState.sortColumn] ?? "");
          const rightValue = String(right[currentState.sortColumn] ?? "");
          const leftNumber = Number(leftValue.replace(/[$,]/g, ""));
          const rightNumber = Number(rightValue.replace(/[$,]/g, ""));
          const comparison = Number.isFinite(leftNumber) && Number.isFinite(rightNumber) && leftValue !== "" && rightValue !== ""
            ? leftNumber - rightNumber
            : leftValue.localeCompare(rightValue, undefined, { numeric: true, sensitivity: "base" });
          return comparison * currentState.sortDirection;
        });
      }
      return rows;
    }

    function exportDataset(datasetId) {
      const dataset = REPORT_DATA[datasetId];
      const rows = getVisibleRows(datasetId);
      const csv = [
        dataset.columns,
        ...rows.map(row => dataset.columns.map(column => row[column] ?? "")),
      ].map(row => row.map(csvCell).join(",")).join("\\r\\n");
      const link = document.createElement("a");
      link.href = URL.createObjectURL(new Blob(["\\uFEFF", csv], { type: "text/csv;charset=utf-8" }));
      link.download = exportFilename(dataset);
      document.body.appendChild(link);
      link.click();
      link.remove();
      URL.revokeObjectURL(link.href);
    }

    function renderSummary(panel) {
      const ruleSummary = REPORT_META.ruleSummary || {};
      const ruleText = [
        `Authorization: ${ruleSummary.authorizationActivity || "any activity"}`,
        `Batch evidence: ${ruleSummary.batchEvidence || "any batch record"}`,
        `Last-batch lookback: ${ruleSummary.lastBatchLookbackDays || "-"} day(s)`,
      ].join(" | ");
      const topRows = [...REPORT_DATA.primary.rows]
        .sort((left, right) => Number(String(right.AMOUNT || right.approvedAmount || "0").replace(/[$,]/g, "")) - Number(String(left.AMOUNT || left.approvedAmount || "0").replace(/[$,]/g, "")))
        .slice(0, 5);
      const topRowsHtml = topRows.length
        ? topRows.map(row => `<tr><td>${escapeHtml(row.STORENAME || "(unnamed store)")}</td><td>${escapeHtml(row.accountNumber || "-")}</td><td>${escapeHtml(row.lastBatchDate || "No history")}</td><td class="amount">${formatMoney(row.AMOUNT || row.approvedAmount)}</td></tr>`).join("")
        : `<tr><td colspan="4" class="empty">No primary exception rows were generated.</td></tr>`;
      const filesHtml = TAB_DEFINITIONS.filter(tab => tab.id !== "summary" && REPORT_DATA[tab.id]).map(tab => `<li><span>${escapeHtml(REPORT_DATA[tab.id].filename)}</span><small>${formatNumber(REPORT_DATA[tab.id].rows.length)} row(s)</small></li>`).join("");
      panel.innerHTML = `
        <section class="finding"><strong>Primary finding.</strong> ${escapeHtml(REPORT_META.finding)}</section>
        <h2>Key results</h2>
        <section class="cards">
          <div class="card"><div class="label">Stores flagged</div><div class="value">${formatNumber(REPORT_META.primaryStoreCount)}</div></div>
          <div class="card green"><div class="label">Authorized-unbatched</div><div class="value">${formatMoney(REPORT_META.primaryTotal)}</div></div>
          <div class="card"><div class="label">Primary rows</div><div class="value">${formatNumber(REPORT_META.primaryRowCount)}</div></div>
          <div class="card yellow"><div class="label">Batch records</div><div class="value">${formatNumber(REPORT_META.batchRecordCount)}</div></div>
          <div class="card red"><div class="label">Review excluded</div><div class="value">${formatNumber(REPORT_META.reviewCount)}</div></div>
        </section>
        <div class="two-col">
          <section><h2>Top flagged accounts</h2><div class="panel"><div class="panel-title">Authorized-unbatched exposure</div><table><thead><tr><th>Store</th><th>Account</th><th>Last batch</th><th class="amount">Amount</th></tr></thead><tbody>${topRowsHtml}</tbody></table></div></section>
          <section><h2>Embedded data</h2><div class="panel"><ul class="files">${filesHtml}</ul></div></section>
        </div>
        <h2>Important interpretation</h2>
        <div class="note">The primary report is account/term level. termID values are linked through accountNumber to batch history; lastBatchDate is the latest batch timestamp in that history. Authorized-unbatched amounts are account-level approved exposure, not proof that a specific physical terminal failed to batch. Rule set: ${escapeHtml(ruleText)}. Review exclusions by reason: ${escapeHtml(REPORT_META.reasonText)}. ${escapeHtml(REPORT_META.ambiguousText)}</div>`;
    }

    function renderTable(panel, datasetId) {
      const dataset = REPORT_DATA[datasetId];
      const currentState = state[datasetId];
      const rows = getVisibleRows(datasetId);
      panel.querySelector(".row-count").textContent = `${formatNumber(rows.length)} of ${formatNumber(dataset.rows.length)} rows`;
      const body = panel.querySelector("tbody");
      body.replaceChildren();
      if (!rows.length) {
        const emptyRow = document.createElement("tr");
        const emptyCell = document.createElement("td");
        emptyCell.colSpan = dataset.columns.length;
        emptyCell.className = "empty";
        emptyCell.textContent = "No matching rows.";
        emptyRow.appendChild(emptyCell);
        body.appendChild(emptyRow);
        return;
      }
      rows.forEach(row => {
        const tableRow = document.createElement("tr");
        dataset.columns.forEach(column => {
          const cell = document.createElement("td");
          cell.textContent = row[column] ?? "";
          tableRow.appendChild(cell);
        });
        body.appendChild(tableRow);
      });
    }

    function renderDataset(panel, datasetId) {
      const dataset = REPORT_DATA[datasetId];
      state[datasetId] = state[datasetId] || { search: "", sortColumn: null, sortDirection: 1 };
      panel.innerHTML = `<div class="toolbar"><input class="search" type="search" placeholder="Search this tab..."><button class="export-button" type="button">Export CSV</button><span class="row-count"></span></div><div class="table-wrap"><table><thead><tr></tr></thead><tbody></tbody></table></div>`;
      const headerRow = panel.querySelector("thead tr");
      dataset.columns.forEach(column => {
        const header = document.createElement("th");
        header.textContent = column;
        header.title = "Click to sort";
        header.addEventListener("click", () => {
          const currentState = state[datasetId];
          currentState.sortDirection = currentState.sortColumn === column ? currentState.sortDirection * -1 : 1;
          currentState.sortColumn = column;
          renderTable(panel, datasetId);
        });
        headerRow.appendChild(header);
      });
      const search = panel.querySelector(".search");
      search.value = state[datasetId].search;
      search.addEventListener("input", event => {
        state[datasetId].search = event.target.value;
        renderTable(panel, datasetId);
      });
      panel.querySelector(".export-button").addEventListener("click", () => exportDataset(datasetId));
      renderTable(panel, datasetId);
    }

    function activateTab(tabId) {
      document.querySelectorAll(".tab").forEach(button => button.classList.toggle("active", button.dataset.tab === tabId));
      const panel = document.getElementById(`panel-${tabId}`);
      document.querySelectorAll(".tab-panel").forEach(section => section.hidden = section !== panel);
      if (tabId === "summary") renderSummary(panel);
      else renderDataset(panel, tabId);
    }

    TAB_DEFINITIONS.forEach((tab, index) => {
      const button = document.createElement("button");
      button.className = "tab";
      button.dataset.tab = tab.id;
      button.type = "button";
      button.textContent = tab.label;
      button.addEventListener("click", () => activateTab(tab.id));
      tabsElement.appendChild(button);
      const panel = document.createElement("section");
      panel.id = `panel-${tab.id}`;
      panel.className = "tab-panel";
      panel.hidden = index !== 0;
      contentElement.appendChild(panel);
    });
    activateTab("summary");
  </script>
</body>
</html>
"""
    html_document = (
        html_document
        .replace("__LOGO__", logo_html)
        .replace("__GENERATED__", escape(generated))
        .replace("__BATCH_START__", escape(report_start.strftime("%Y-%m-%d")))
        .replace("__BATCH_END__", escape(report_end.strftime("%Y-%m-%d")))
        .replace("__AUTH_WINDOW__", escape(auth_window))
        .replace("__REPORT_DATA__", report_data_json)
        .replace("__REPORT_META__", report_meta_json)
    )
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(html_document, encoding="utf-8")
    return path


def make_client(config):
    base_url = text(config.get("apiBaseUrl")) or BASE_URL
    if base_url.endswith(AUTH_PATH):
        base_url = base_url[: -len(AUTH_PATH)].rstrip("/")
    api_key_name = text(config.get("apiKeyEnvironmentVariable")) or "MXCONNECT_API_KEY"
    api_key = clean_api_key(EMBEDDED_API_KEY) or clean_api_key(os.environ.get(api_key_name))
    if not api_key:
        raise RuntimeError(
            f"Set {api_key_name} or use an executable built with an embedded API key."
        )
    return MxConnectClient(base_url, api_key)


def parse_args():
    parser = ArgumentParser(description="Generate the store-level TSYS PAX batch report.")
    parser.add_argument(
        "--ui", action="store_true",
        help="Open the configuration UI (default when no command is supplied)",
    )
    parser.add_argument(
        "--run-report", action="store_true",
        help="Run the report without opening the UI; intended for Task Scheduler",
    )
    parser.add_argument(
        "-c", "--config", default=None,
        help="JSON configuration file (default: ./config.json)",
    )
    parser.add_argument(
        "--batch-days", type=int, default=None,
        help="Batch lookback in days (default: config batchLookbackDays or 3)",
    )
    parser.add_argument(
        "--historical-days", type=int, default=None,
        help="Historical batch export lookback (default: config lastBatchLookbackDays or 60)",
    )
    parser.add_argument(
        "--auth-check-days", type=int, default=None,
        help="Absolute authorization activity window (default: batch-days)",
    )
    parser.add_argument(
        "--auth-window", default=None,
        help="Optional legacy MXConnect quick authorization window override",
    )
    parser.add_argument(
        "--log-file", default=None,
        help="Internal log file path used when the windowed executable runs in report mode",
    )
    return parser.parse_args()


def main():
    started = perf_counter()
    args = parse_args()
    configure_process_output(args.log_file)

    config_path = Path(args.config or "config.json").expanduser().resolve()
    config, raw_devices = load_config(config_path)

    batch_days = args.batch_days or int(config.get("batchLookbackDays", 3))
    historical_days = max(
        batch_days,
        args.historical_days
        or int(config.get("lastBatchLookbackDays", config.get("historicalLookbackDays", DEFAULT_LAST_BATCH_LOOKBACK_DAYS))),
    )
    auth_check_days = max(
        1,
        getattr(args, "auth_check_days", None)
        or int(config.get("authCheckDays", batch_days or DEFAULT_AUTH_CHECK_DAYS)),
    )
    configured_auth_mode = text(config.get("authorizationMode")).casefold() or "absolute"
    auth_window_override = getattr(args, "auth_window", None)
    if not auth_window_override and configured_auth_mode == "quick":
        auth_window_override = text(config.get("authorizationWindow")) or "last_24_h"
    base_url = text(config.get("apiBaseUrl")) or BASE_URL
    if base_url.endswith(AUTH_PATH):
        base_url = base_url[: -len(AUTH_PATH)].rstrip("/")

    output_directory = create_timestamped_output_directory(
        resolve_output_directory(config_path, config)
    )
    report_catalog = AccountReportCatalog(
        output_directory.parent / REPORT_CATALOG_FILENAME
    )
    if report_catalog.valid and report_catalog.path.exists():
        print(f"Validated report catalog: {report_catalog.path}")
    elif not report_catalog.valid:
        print(f"Warning: {report_catalog.validation_message}")

    print(f"Loaded {len(raw_devices)} optional store display overrides.")
    print(f"MXConnect base URL: {base_url}")
    print(
        f"Batch window: {batch_days} days; authorization window: "
        f"{auth_window_override or f'absolute {auth_check_days} days (UTC)'}; "
        f"last-batch lookback: {historical_days} days"
    )
    print(f"Run output directory: {output_directory}")
    flush_output()

    print("Step 1 of 5: Connecting to MXConnect...")
    client = make_client(config)
    client.authenticate()

    report_start, report_end = date_window(batch_days)
    print("Step 2 of 5: Checking batch activity...")
    batch_records = fetch_all(
        client,
        batch_url(base_url, report_start.strftime("%Y-%m-%d"), report_end.strftime("%Y-%m-%d")),
    )
    accepted_batch_records = [record for record in batch_records if is_accepted_batch(record)]
    batch_status_counts = Counter(batch_status(record) for record in batch_records)
    print(
        f"Batch activity reviewed: {len(batch_records)} records; "
        f"{batch_status_counts.get('accepted', 0)} accepted, "
        f"{batch_status_counts.get('rejected', 0)} rejected, "
        f"{batch_status_counts.get('unknown', 0)} unknown"
    )
    print("Step 3 of 5: Loading active TSYS merchants...")
    roster_records = fetch_uar(client)
    active_accounts = active_tsys_accounts(roster_records)
    print(f"Active TSYS merchants: {len(active_accounts)}")

    current_records = batch_records
    batched_accounts = {
        text(record.get("accountNumber"))
        for record in current_records
        if text(record.get("accountNumber"))
    }
    print(f"Accounts with a batch record in the report window: {len(batched_accounts)}")

    candidate_accounts = active_accounts - batched_accounts
    print(f"Active TSYS stores with no batch record: {len(candidate_accounts)}")

    history_records_for_catalog = batch_records
    history_floor_date = report_start.strftime("%Y-%m-%d")
    if candidate_accounts and historical_days > batch_days:
        history_start, history_end = date_window(historical_days)
        history_floor_date = history_start.strftime("%Y-%m-%d")
        print(
            f"Supplementing lastBatchDate from the {historical_days}-day batch history..."
        )
        history_records_for_catalog = fetch_all(
            client,
            batch_url(
                base_url,
                history_start.strftime("%Y-%m-%d"),
                history_end.strftime("%Y-%m-%d"),
            ),
        )
        print(
            f"Historical batch activity reviewed: {len(history_records_for_catalog)} records"
        )
    history_fingerprint = fingerprint_records(history_records_for_catalog)
    cached_history_views = report_catalog.cached_history_views(history_fingerprint)
    if cached_history_views is not None:
        print("Reusing the validated three-report history catalog.")
        history_catalog = BatchHistoryCatalog.from_cached_views(cached_history_views)
    else:
        history_catalog = BatchHistoryCatalog(history_records_for_catalog)
    last_batch_by_account = history_catalog.last_batch_by_account
    term_history_accounts = history_catalog.term_history_by_account

    print("Step 4 of 5: Checking authorization activity...")
    if auth_window_override:
        auth_window = f"quick:{auth_window_override}"
        activity_url = authorization_url(base_url, auth_window_override, candidate_accounts)
    else:
        auth_start, auth_end = date_window(auth_check_days)
        auth_window = (
            f"absolute:{auth_start.strftime('%Y-%m-%d')} to "
            f"{auth_end.strftime('%Y-%m-%d')} UTC"
        )
        activity_url = authorization_absolute_url(
            base_url,
            auth_start.strftime("%Y-%m-%d"),
            auth_end.strftime("%Y-%m-%d"),
            candidate_accounts,
        )
    authorization_records = fetch_all(client, activity_url) if candidate_accounts else []
    activity_summary = build_account_activity_summary(authorization_records)
    in_use_accounts = set(activity_summary)
    print(
        f"Authorization activity reviewed: {len(authorization_records)} records; "
        f"{len(in_use_accounts)} stores active"
    )

    print("Calculating approved authorized-unbatched amounts...")
    today_date = report_end.strftime("%Y-%m-%d")
    authorized_unbatched_by_account = {}
    for account in sorted(in_use_accounts):
        since_date = iso_date_from_display(
            last_batch_by_account.get(account),
            history_floor_date,
        )
        authorized_unbatched_by_account[account] = authorized_unbatched_amount(
            client,
            base_url,
            account,
            since_date,
            today_date,
        )

    email_rows, review_rows = create_store_report(
        roster_records,
        raw_devices,
        activity_summary,
        current_records,
        last_batch_by_account,
        term_history_accounts,
        authorized_unbatched_by_account,
        history_floor_date,
    )

    print("Step 5 of 5: Writing HTML report...")
    summary_path = output_directory / "BottlePOS PAX Batch Report.html"
    term_history_rows = history_catalog.term_history_rows
    batch_history_rows = compact_batch_history_rows(batch_records)
    report_scope = f"{auth_window}; last-batch lookback: {historical_days} days"
    write_interactive_summary_html(
        summary_path,
        batch_days,
        report_scope,
        report_start,
        report_end,
        [
            ("Primary exception report", "PINPAD_BATCH_NOT_CLOSED", "Stores that met the primary report criteria.", email_rows),
            ("Store review", "NEEDS_MAPPING_OR_REVIEW", "Rows excluded from the primary report for review.", review_rows),
            (BATCH_HISTORY_LABEL, "BATCH_HISTORY", "Compact batch history used by this run, including batch-status evidence.", batch_history_rows),
            ("Term/account history", "TERMID_ACCOUNT_HISTORY", "AccountNumber and termID pairs from batch history.", term_history_rows),
        ],
        {
            "authorizationActivity": "any authorization activity",
            "batchEvidence": "any batch record",
            "lastBatchLookbackDays": historical_days,
            "accountNormalization": "BPOS text normalization",
            "batchStatusCounts": dict(batch_status_counts),
        },
    )

    # Only the derived historical term/account view is eligible for reuse.
    # Authorization totals, current batch state, and terminal identity remain live.
    account_observations = {
        account: {
            TERM_HISTORY_CONSTANT_FIELD: term_history_accounts.get(account) or None,
        }
        for account in candidate_accounts
    }
    report_catalog.record_report(
        output_directory.name,
        account_observations,
        history_fingerprint,
        {
            "lastBatchByAccount": last_batch_by_account,
            "termHistoryByAccount": term_history_accounts,
            "termHistoryRows": term_history_rows,
        },
    )
    try:
        report_catalog.save()
    except OSError as error:
        print(f"Warning: report catalog was not saved: {error}")

    print("Report complete.")
    print(f"Stores flagged: {len(email_rows)}")
    print(f"Stores excluded for review: {len(review_rows)}")
    print("Reports created: 1 HTML report")
    print(f"Output directory: {output_directory}")
    print(f"Elapsed time: {perf_counter() - started:.1f} seconds")
    return 0


UI_COLUMNS = [
    ("enabled", "Enabled", 70),
    ("storeName", "Store Name", 220),
    ("url", "URL", 220),
    ("device", "Device", 85),
    ("accountNumber", "Account Number", 160),
]


def default_config():
    return {
        "apiBaseUrl": BASE_URL,
        "apiKeyEnvironmentVariable": "MXCONNECT_API_KEY",
        "batchLookbackDays": 3,
        "lastBatchLookbackDays": 60,
        "authCheckDays": 3,
        "authorizationMode": "absolute",
        "authorizationWindow": "last_24_h",
        "requireAuthorizationActivity": True,
        "outputDirectory": "./tsys-auditdata",
        "devices": [],
    }


class BatchReportUi(tk.Tk):
    def __init__(self, initial_config=None):
        super().__init__()
        self.title("BottlePOS PAX Batch Report")
        self.geometry("1040x720")
        self.minsize(960, 640)
        self.configure(bg="#ffffff")

        self.script_path = application_entrypoint()
        self.config_path_var = tk.StringVar(
            value=str(Path(initial_config).expanduser().resolve())
            if initial_config else str(self.script_path.with_name("config.json"))
        )
        self.output_directory_var = tk.StringVar(value="./tsys-auditdata")
        self.batch_days_var = tk.IntVar(value=3)
        self.devices = []
        self.config = default_config()
        self.loaded_config_path = None
        self.process = None
        self.edit_control = None
        self.active_command = None
        self.last_output_directory = None

        self.configure_styles()
        self.build_controls()
        self.load_config_path(self.config_path_var.get(), log=False)

    def configure_styles(self):
        style = ttk.Style(self)
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass
        surface = "#ffffff"
        ink = "#1f2328"
        accent = "#0969da"
        control_surface = "#f6f8fa"
        border = "#d0d7de"
        selected = "#ddf4ff"

        style.configure(".", font=("Segoe UI", 9), background=surface, foreground=ink)
        style.configure("TFrame", background=surface)
        style.configure("TLabel", background=surface, foreground=ink)
        style.configure("Section.TLabel", background=surface, foreground="#203239", font=("Segoe UI", 14, "bold"))
        style.configure("Muted.TLabel", background=surface, foreground="#7a8790", font=("Segoe UI", 9))
        style.configure(
            "TButton",
            background=control_surface,
            foreground=ink,
            bordercolor=border,
            lightcolor=border,
            darkcolor=border,
            padding=(10, 6),
            relief="flat",
        )
        style.map(
            "TButton",
            background=[("active", selected), ("pressed", "#b6e3ff"), ("disabled", "#f6f8fa")],
            foreground=[("disabled", "#8c959f")],
            bordercolor=[("focus", accent), ("active", accent)],
        )
        style.configure(
            "TEntry",
            fieldbackground=surface,
            foreground=ink,
            bordercolor=border,
            lightcolor=border,
            darkcolor=border,
            insertcolor=accent,
            padding=(6, 5),
        )
        style.map(
            "TEntry",
            bordercolor=[("focus", accent)],
            lightcolor=[("focus", accent)],
            darkcolor=[("focus", accent)],
        )
        style.configure(
            "Treeview",
            background=surface,
            foreground=ink,
            fieldbackground=surface,
            bordercolor=border,
            lightcolor=border,
            darkcolor=border,
            rowheight=26,
        )
        style.configure(
            "Treeview.Heading",
            background=control_surface,
            foreground=ink,
            bordercolor=border,
            lightcolor=border,
            darkcolor=border,
            padding=(8, 6),
            relief="flat",
        )
        style.map(
            "Treeview",
            background=[("selected", selected)],
            foreground=[("selected", ink)],
        )
        style.configure(
            "TScrollbar",
            background=control_surface,
            troughcolor=surface,
            bordercolor=border,
            arrowcolor=ink,
        )
        style.map("TScrollbar", background=[("active", "#d8dee4")])

    def build_controls(self):
        shell = tk.Frame(self, bg="#f4f5f6")
        shell.pack(fill="both", expand=True)
        content = tk.Frame(shell, bg="#f4f5f6")
        content.pack(fill="both", expand=True)
        header = tk.Frame(content, height=66, bg="#ffffff", highlightthickness=1, highlightbackground="#e1e4e6")
        header.pack(fill="x")
        header.pack_propagate(False)
        logo_root = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
        logo_path = logo_root / "bottlepos_logo.png"
        self.logo_image = None
        if logo_path.exists():
            try:
                self.logo_image = tk.PhotoImage(file=str(logo_path))
            except tk.TclError:
                self.logo_image = None
        if self.logo_image is not None:
            tk.Label(header, image=self.logo_image, bg="#ffffff").pack(side="left", padx=(18, 12))
        tk.Label(
            header,
            text="PAX BATCH REPORT",
            bg="#ffffff",
            fg="#203239",
            font=("Segoe UI", 15, "bold"),
            anchor="w",
        ).pack(side="left")
        tk.Label(
            header,
            text="Bottle POS",
            bg="#ffffff",
            fg="#3d8fca",
            font=("Segoe UI", 10, "bold"),
            anchor="e",
        ).pack(side="right", padx=20)

        outer = ttk.Frame(content, padding=16)
        outer.pack(fill="both", expand=True)
        ttk.Label(outer, text="Report setup", style="Section.TLabel").pack(anchor="w", pady=(0, 8))

        top = ttk.Frame(outer)
        top.pack(fill="x")
        top.columnconfigure(1, weight=1)

        ttk.Label(top, text="Config file").grid(row=0, column=0, sticky="w", padx=(0, 8), pady=4)
        ttk.Entry(top, textvariable=self.config_path_var).grid(row=0, column=1, sticky="ew", pady=4)
        ttk.Button(top, text="Open config", command=self.open_config).grid(row=0, column=2, padx=(8, 0), pady=4)

        ttk.Label(top, text="Output folder").grid(row=1, column=0, sticky="w", padx=(0, 8), pady=4)
        ttk.Entry(top, textvariable=self.output_directory_var).grid(row=1, column=1, sticky="ew", pady=4)
        ttk.Button(top, text="Browse folder", command=self.browse_output).grid(row=1, column=2, padx=(8, 0), pady=4)

        options = ttk.Frame(top)
        options.grid(row=2, column=0, columnspan=3, sticky="w", pady=(6, 4))
        ttk.Label(options, text="Report timeframe (days)").pack(side="left", padx=(0, 5))
        tk.Spinbox(
            options,
            from_=1,
            to=365,
            width=5,
            textvariable=self.batch_days_var,
            bg="#ffffff",
            fg="#1f2328",
            insertbackground="#0969da",
            buttonbackground="#f6f8fa",
            relief="solid",
            bd=1,
            highlightthickness=1,
            highlightbackground="#d0d7de",
            highlightcolor="#0969da",
        ).pack(side="left", padx=(0, 14))

        actions = ttk.Frame(top)
        actions.grid(row=3, column=0, columnspan=3, sticky="w", pady=(4, 8))
        ttk.Button(actions, text="Save config", command=self.save_config_from_ui).pack(side="left", padx=(0, 8))
        ttk.Button(actions, text="Run report", command=self.run_report).pack(side="left", padx=(0, 8))

        self.log_box = tk.Text(
            outer,
            height=12,
            bg="#f6f8fa",
            fg="#1f2328",
            insertbackground="#0969da",
            selectbackground="#ddf4ff",
            selectforeground="#1f2328",
            relief="solid",
            bd=1,
            highlightthickness=1,
            highlightbackground="#d0d7de",
            highlightcolor="#0969da",
            wrap="word",
        )
        self.log_box.pack(fill="both", expand=True, pady=(16, 0))

    def log(self, message):
        self.log_box.insert("end", f"[{datetime.now():%H:%M:%S}] {message}\n")
        self.log_box.see("end")

    def load_config_path(self, raw_path, log=True):
        path = Path(raw_path).expanduser().resolve()
        self.config_path_var.set(str(path))
        if not path.exists():
            self.config = default_config()
            self.devices = []
            self.loaded_config_path = str(path)
            self.output_directory_var.set(str((path.parent / "tsys-auditdata").resolve()))
            self.refresh_table()
            if log:
                self.log(f"Config does not exist yet: {path}")
            return
        try:
            config, devices = load_config(path)
            self.config = config
            self.devices = [dict(device) for device in devices]
            self.loaded_config_path = str(path)
            self.output_directory_var.set(str(resolve_output_directory(path, config)))
            self.batch_days_var.set(int(config.get("batchLookbackDays", 3)))
            self.refresh_table()
            if log:
                self.log(f"Loaded {len(self.devices)} devices.")
        except Exception as error:
            self.log(f"ERROR: {error}")

    def open_config(self):
        path = filedialog.askopenfilename(
            title="Open config",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
        )
        if path:
            self.load_config_path(path)

    def browse_output(self):
        path = filedialog.askdirectory(
            title="Select output folder",
            initialdir=self.output_directory_var.get() or str(Path.cwd()),
        )
        if path:
            self.output_directory_var.set(path)

    def display_value(self, device, key):
        if key == "enabled":
            return "Yes" if device.get(key, True) else "No"
        return text(device.get(key))

    def refresh_table(self):
        # The setup UI no longer renders the optional device override table.
        # The loaded device rows are preserved and saved back to config.json.
        return

    def begin_cell_edit(self, event):
        row_id = self.tree.identify_row(event.y)
        column_id = self.tree.identify_column(event.x)
        if not row_id or not column_id:
            return
        column_index = int(column_id[1:]) - 1
        if column_index < 0 or column_index >= len(UI_COLUMNS):
            return
        device_index = int(row_id)
        key = UI_COLUMNS[column_index][0]
        if key == "enabled":
            self.devices[device_index][key] = not bool(self.devices[device_index].get(key, True))
            self.refresh_table()
            return
        bounds = self.tree.bbox(row_id, column_id)
        if not bounds:
            return
        x, y, width, height = bounds
        self.commit_cell_edit()
        self.edit_control = tk.Entry(
            self.tree,
            bg="#ffffff",
            fg="#1f2328",
            insertbackground="#0969da",
            selectbackground="#ddf4ff",
            selectforeground="#1f2328",
            relief="solid",
            bd=1,
            highlightthickness=1,
            highlightbackground="#0969da",
            highlightcolor="#0969da",
        )
        self.edit_control.insert(0, text(self.devices[device_index].get(key)))
        self.edit_control.place(x=x, y=y, width=width, height=height)
        self.edit_control.focus_set()
        self.edit_control.bind("<Return>", lambda _: self.commit_cell_edit(device_index, key))
        self.edit_control.bind("<Escape>", lambda _: self.commit_cell_edit())
        self.edit_control.bind("<FocusOut>", lambda _: self.commit_cell_edit(device_index, key))

    def commit_cell_edit(self, device_index=None, key=None):
        if self.edit_control is None:
            return
        if device_index is not None and key is not None:
            self.devices[device_index][key] = self.edit_control.get().strip()
        self.edit_control.destroy()
        self.edit_control = None
        self.refresh_table()

    def save_config_from_ui(self, announce=True):
        try:
            self.commit_cell_edit()
            path = Path(self.config_path_var.get()).expanduser().resolve()
            config = dict(self.config or default_config())
            config["batchLookbackDays"] = max(1, int(self.batch_days_var.get()))
            config["historicalLookbackDays"] = max(
                int(config.get("historicalLookbackDays", 90)),
                config["batchLookbackDays"],
            )
            config["requireAuthorizationActivity"] = True
            save_config(path, config, self.devices, self.output_directory_var.get())
            self.config = config
            self.loaded_config_path = str(path)
            self.config_path_var.set(str(path))
            if announce:
                self.log(f"Saved {path}.")
            return path
        except Exception as error:
            self.log(f"ERROR: {error}")
            return None

    def start_command(self, command):
        if self.process is not None and self.process.poll() is None:
            self.log("A report is already running.")
            return
        config_path = self.save_config_from_ui()
        if config_path is None:
            return
        self.active_command = command[0] if command else None
        self.last_output_directory = None
        log_path = Path(tempfile.gettempdir()) / (
            f"TSYS_PAX_BATCH_REPORT_{os.getpid()}_{datetime.now():%Y%m%d%H%M%S%f}.log"
        )
        command = (
            [str(self.script_path), *command, "--config", str(config_path), "--log-file", str(log_path)]
            if getattr(sys, "frozen", False)
            else [sys.executable, str(self.script_path), *command, "--config", str(config_path), "--log-file", str(log_path)]
        )
        try:
            self.process = subprocess.Popen(
                command,
                cwd=str(self.script_path.parent),
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                env=os.environ.copy(),
            )
            threading.Thread(
                target=self.read_process_output,
                args=(self.process, log_path),
                daemon=True,
            ).start()
        except Exception as error:
            self.active_command = None
            self.log(f"ERROR: {error}")

    def read_process_output(self, process, log_path):
        position = 0

        def read_new_lines():
            nonlocal position
            if not log_path.exists():
                return
            with log_path.open("r", encoding="utf-8", errors="replace") as log_file:
                log_file.seek(position)
                lines = log_file.readlines()
                position = log_file.tell()
            for line in lines:
                self.after(0, self.handle_process_line, line.rstrip())

        while process.poll() is None:
            read_new_lines()
            threading.Event().wait(0.1)
        read_new_lines()
        exit_code = process.wait()
        try:
            log_path.unlink()
        except OSError:
            pass
        self.after(0, self.report_finished, exit_code)

    def report_finished(self, exit_code):
        self.log("Finished." if exit_code == 0 else f"ERROR: report exited with code {exit_code}.")
        self.process = None
        if exit_code == 0 and self.active_command == "--run-report":
            if self.last_output_directory is None:
                self.log("ERROR: Report completed, but the output directory was not reported.")
            else:
                self.show_completion_popup(
                    self.last_output_directory / "BottlePOS PAX Batch Report.html"
                )
        self.active_command = None

    def handle_process_line(self, line):
        self.log(line)
        if line.startswith("Output directory:"):
            self.last_output_directory = Path(line.partition(":")[2].strip())

    def show_completion_popup(self, report_path):
        popup = tk.Toplevel(self)
        popup.title("Report complete")
        popup.configure(bg="#ffffff")
        popup.resizable(False, False)
        popup.transient(self)
        popup.protocol("WM_DELETE_WINDOW", self.destroy)

        body = tk.Frame(popup, bg="#ffffff", padx=24, pady=20)
        body.pack(fill="both", expand=True)
        tk.Label(
            body,
            text="Report generation complete",
            bg="#ffffff",
            fg="#203239",
            font=("Segoe UI", 13, "bold"),
        ).pack(anchor="w")
        tk.Label(
            body,
            text="The interactive HTML report is ready to view.",
            bg="#ffffff",
            fg="#7a8790",
            font=("Segoe UI", 9),
        ).pack(anchor="w", pady=(5, 16))

        buttons = tk.Frame(body, bg="#ffffff")
        buttons.pack(fill="x")

        def view_report():
            if not report_path.exists():
                self.log(f"ERROR: Report file was not found: {report_path}")
                return
            try:
                webbrowser.open(report_path.resolve().as_uri())
                popup.destroy()
            except OSError as error:
                self.log(f"ERROR: Could not open report: {error}")

        tk.Button(
            buttons,
            text="VIEW REPORT",
            command=view_report,
            bg="#0969da",
            fg="#ffffff",
            activebackground="#0550ae",
            activeforeground="#ffffff",
            relief="flat",
            borderwidth=0,
            padx=14,
            pady=8,
            cursor="hand2",
        ).pack(side="left", padx=(0, 8))
        tk.Button(
            buttons,
            text="CLOSE",
            command=self.destroy,
            bg="#f6f8fa",
            fg="#1f2328",
            activebackground="#ddf4ff",
            activeforeground="#1f2328",
            relief="flat",
            borderwidth=1,
            padx=14,
            pady=8,
            cursor="hand2",
        ).pack(side="left")

        popup.update_idletasks()
        x = self.winfo_rootx() + max(0, (self.winfo_width() - popup.winfo_width()) // 2)
        y = self.winfo_rooty() + max(0, (self.winfo_height() - popup.winfo_height()) // 2)
        popup.geometry(f"+{x}+{y}")
        popup.grab_set()
        popup.focus_set()

    def run_report(self):
        self.start_command(["--run-report"])


def launch_ui(initial_config=None):
    BatchReportUi(initial_config).mainloop()


if __name__ == "__main__":
    try:
        if len(sys.argv) == 1 or "--ui" in sys.argv:
            initial_config = None
            for flag in ("--config", "-c"):
                if flag in sys.argv:
                    position = sys.argv.index(flag)
                    if position + 1 < len(sys.argv):
                        initial_config = sys.argv[position + 1]
                    break
            launch_ui(initial_config)
        else:
            raise SystemExit(main())
    except (requests.RequestException, RuntimeError, ValueError, OSError, json.JSONDecodeError) as error:
        print(f"ERROR: {error}", file=sys.stderr)
        raise SystemExit(1)
